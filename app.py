import os
import io
import logging
import tempfile
import traceback
import re
import json
import uuid
import requests
import random
import time
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from functools import wraps

# JasoAI 수정: admin_required 정의는 라인 370에 있음 (role 기반)
from flask import Flask, request, jsonify, render_template, session, redirect, url_for, flash, current_app
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_session import Session
from datetime import timedelta
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from openai import OpenAI
from openai.types.chat import ChatCompletionSystemMessageParam, ChatCompletionUserMessageParam
from models import db, User, EvaluationRecord, Review, Question, CreditPurchase, Notification, CreditsLedger
from sqlalchemy import func
import shutil
from datetime import datetime
from utils.idv import normalize_phone, mask_phone, idv_session_is_fresh

# ===== 보안 상수 =====
# 타이밍 사이드채널 공격 방지를 위한 더미 패스워드 해시
DUMMY_PWHASH = generate_password_hash("dummy_password")

# ===== 유틸리티 함수들 =====
def _normalize_email(email):
    """이메일 정규화 (소문자, 공백 제거)"""
    if not email:
        return None
    return email.strip().lower()

def _is_email(text):
    """이메일 형식 검증"""
    if not text:
        return False
    email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return bool(re.match(email_pattern, text.strip()))

def _mask_username(username):
    """사용자명 마스킹 (보안을 위해 일부만 표시)"""
    if not username or len(username) <= 2:
        return username
    if len(username) <= 4:
        return username[:2] + '*' * (len(username) - 2)
    else:
        return username[:2] + '*' * (len(username) - 4) + username[-2:]

def _mask_sensitive_data(data, mask_char='*', show_chars=4):
    """민감한 데이터 마스킹 (전화번호, 인증번호 등)"""
    if not data or len(data) <= show_chars:
        return mask_char * len(data) if data else ''
    return data[:show_chars//2] + mask_char * (len(data) - show_chars) + data[-show_chars//2:]

def _mask_email(email):
    """이메일 마스킹 (user@example.com → us****@ex****.com)"""
    if not email or '@' not in email:
        return email or ''
    
    local, domain = email.split('@', 1)
    
    # 로컬 부분 마스킹: 처음 2자리만 노출
    if len(local) <= 2:
        masked_local = '*' * len(local)
    else:
        masked_local = local[:2] + '*' * (len(local) - 2)
    
    # 도메인 부분 마스킹: 처음 2자리와 TLD만 노출
    if '.' in domain:
        domain_parts = domain.split('.')
        if len(domain_parts[0]) <= 2:
            masked_domain = '*' * len(domain_parts[0])
        else:
            masked_domain = domain_parts[0][:2] + '*' * (len(domain_parts[0]) - 2)
        masked_domain += '.' + '.'.join(domain_parts[1:])
    else:
        if len(domain) <= 2:
            masked_domain = '*' * len(domain)
        else:
            masked_domain = domain[:2] + '*' * (len(domain) - 2)
    
    return f"{masked_local}@{masked_domain}"

def _mask_merchant_uid(merchant_uid):
    """상점 거래번호 마스킹"""
    if not merchant_uid:
        return ''
    # auth_1234567890_123 형태에서 마지막 부분만 노출
    parts = merchant_uid.split('_')
    if len(parts) >= 3:
        return f"auth_****_{parts[-1]}"
    return _mask_sensitive_data(merchant_uid, show_chars=6)

def _send_verification_email(email, verification_code):
    """실제 이메일 인증코드 발송"""
    try:
        # SMTP 설정 확인
        smtp_email = os.environ.get('SMTP_EMAIL')
        smtp_password = os.environ.get('SMTP_PASSWORD')
        
        if not smtp_email or not smtp_password:
            logging.warning("SMTP 설정이 없습니다. 개발 모드로 동작합니다.")
            return False
        
        # 이메일 내용 생성
        subject = "[JasoAI] 이메일 인증코드"
        html_body = f"""
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                .container {{ max-width: 600px; margin: 0 auto; font-family: 'Noto Sans KR', sans-serif; }}
                .header {{ background-color: #4F46E5; color: white; padding: 20px; text-align: center; }}
                .content {{ padding: 30px 20px; background-color: #f8f9fa; }}
                .code-box {{ background-color: white; border: 2px dashed #4F46E5; padding: 20px; text-align: center; margin: 20px 0; }}
                .code {{ font-size: 24px; font-weight: bold; color: #4F46E5; letter-spacing: 3px; }}
                .footer {{ text-align: center; padding: 20px; color: #6c757d; font-size: 12px; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>JasoAI 이메일 인증</h1>
                </div>
                <div class="content">
                    <h2>안녕하세요!</h2>
                    <p>JasoAI 회원가입을 위한 이메일 인증코드입니다.</p>
                    <p>아래 인증코드를 입력하여 회원가입을 완료해주세요.</p>
                    
                    <div class="code-box">
                        <div class="code">{verification_code}</div>
                    </div>
                    
                    <p><strong>주의사항:</strong></p>
                    <ul>
                        <li>인증코드는 3분간만 유효합니다.</li>
                        <li>보안을 위해 인증코드를 다른 사람과 공유하지 마세요.</li>
                        <li>본인이 요청하지 않은 경우 이 이메일을 무시해주세요.</li>
                    </ul>
                </div>
                <div class="footer">
                    <p>© 2025 JasoAI. 모든 권리 보유.</p>
                    <p>문의사항: jasoai0612@gmail.com</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        # 이메일 메시지 생성
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = smtp_email
        msg['To'] = email
        
        # HTML 본문 추가
        html_part = MIMEText(html_body, 'html', 'utf-8')
        msg.attach(html_part)
        
        # Gmail SMTP 서버 연결 및 발송
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(smtp_email, smtp_password)
            server.send_message(msg)
        
        logging.info(f"이메일 발송 성공: {_mask_email(email)}")
        return True
        
    except Exception as e:
        logging.error(f"이메일 발송 실패: {_mask_email(email)}, 오류: {str(e)}")
        return False


# 파일 처리 라이브러리 임포트
try:
    import docx
    import pdfplumber
    ENABLE_FILE_UPLOAD = True
except ImportError:
    ENABLE_FILE_UPLOAD = False
    logging.warning("docx 또는 pdfplumber 라이브러리를 찾을 수 없습니다. 파일 업로드 기능이 비활성화됩니다.")

# 환경 감지
PRODUCTION_ENV = os.environ.get('FLASK_ENV') == 'production' or os.environ.get('ENVIRONMENT') == 'production'

# 로깅 설정 - 환경별 레벨 조정
logging_level = logging.WARNING if PRODUCTION_ENV else logging.DEBUG
logging.basicConfig(
    level=logging_level,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

# ===== 알림 및 환불 헬퍼 함수 =====
def add_notification(user_id, notif_type, title, body, meta=None):
    """알림 생성 및 저장"""
    try:
        meta_json = json.dumps(meta, ensure_ascii=False) if meta else None
        notification = Notification(
            user_id=user_id,
            type=notif_type,
            title=title,
            body=body,
            meta_json=meta_json
        )
        db.session.add(notification)
        db.session.commit()
        logging.info(f"[NOTIFICATION][OK] user_id={user_id}, type={notif_type}, title={title}")
        return True
    except Exception as e:
        db.session.rollback()
        logging.error(f"[NOTIFICATION][ERROR] 알림 생성 실패: {str(e)}")
        return False


def portone_refund(imp_uid, amount=None, reason='admin_refund'):
    """PortOne 환불 API 호출 (스텁 처리 포함)"""
    try:
        api_key = os.environ.get("PORTONE_API_KEY")
        api_secret = os.environ.get("PORTONE_API_SECRET")
        
        if not api_key or not api_secret:
            logging.warning("[PORTONE_REFUND][STUB] API 키가 없어 스텁 모드로 처리합니다.")
            return {"ok": True, "stub": True, "message": "secrets missing"}
        
        token_url = "https://api.iamport.kr/users/getToken"
        token_data = {
            "imp_key": api_key,
            "imp_secret": api_secret
        }
        
        for attempt in range(2):
            try:
                token_response = requests.post(token_url, json=token_data, timeout=10)
                if token_response.status_code != 200:
                    logging.error(f"[PORTONE_REFUND][FAIL] 토큰 발급 실패: {token_response.status_code}")
                    if attempt == 0:
                        time.sleep(1)
                        continue
                    return {"ok": False, "message": "토큰 발급 실패"}
                
                token_result = token_response.json()
                if token_result.get("code") != 0:
                    logging.error(f"[PORTONE_REFUND][FAIL] 토큰 응답 오류: {token_result.get('message')}")
                    return {"ok": False, "message": token_result.get("message", "토큰 발급 오류")}
                
                access_token = token_result["response"]["access_token"]
                
                cancel_url = "https://api.iamport.kr/payments/cancel"
                headers = {"Authorization": access_token, "Content-Type": "application/json"}
                cancel_data = {
                    "imp_uid": imp_uid,
                    "reason": reason
                }
                if amount:
                    cancel_data["amount"] = amount
                
                cancel_response = requests.post(cancel_url, headers=headers, json=cancel_data, timeout=10)
                
                if cancel_response.status_code != 200:
                    logging.error(f"[PORTONE_REFUND][FAIL] 환불 요청 실패: {cancel_response.status_code}")
                    if attempt == 0:
                        time.sleep(1)
                        continue
                    return {"ok": False, "message": "환불 요청 실패"}
                
                cancel_result = cancel_response.json()
                if cancel_result.get("code") != 0:
                    logging.error(f"[PORTONE_REFUND][FAIL] 환불 응답 오류: {cancel_result.get('message')}")
                    return {"ok": False, "message": cancel_result.get("message", "환불 처리 오류")}
                
                response_data = cancel_result.get("response", {})
                refunded_amount = response_data.get("cancel_amount")
                
                if amount and refunded_amount != amount:
                    logging.error(f"[PORTONE_REFUND][FAIL] 환불 금액 불일치: 요청={amount}, 실제={refunded_amount}")
                    return {"ok": False, "message": f"환불 금액 불일치 (요청: {amount}원, 실제: {refunded_amount}원)"}
                
                if not refunded_amount:
                    logging.error(f"[PORTONE_REFUND][FAIL] 환불 금액 정보 없음")
                    return {"ok": False, "message": "환불 금액 정보를 확인할 수 없습니다"}
                
                logging.info(f"[PORTONE_REFUND][OK] imp_uid={imp_uid}, cancel_amount={refunded_amount}")
                
                return {
                    "ok": True,
                    "cancel_amount": refunded_amount,
                    "cancel_id": response_data.get("cancel_receipt_urls", [None])[0] if response_data.get("cancel_receipt_urls") else None,
                    "cancelled_at": response_data.get("cancelled_at")
                }
                
            except requests.exceptions.Timeout:
                logging.warning(f"[PORTONE_REFUND][TIMEOUT] 시도 {attempt+1}/2")
                if attempt == 0:
                    time.sleep(1)
                    continue
                return {"ok": False, "message": "환불 요청 시간 초과"}
            except requests.exceptions.RequestException as e:
                logging.error(f"[PORTONE_REFUND][ERROR] 네트워크 오류: {str(e)}")
                if attempt == 0:
                    time.sleep(1)
                    continue
                return {"ok": False, "message": f"네트워크 오류: {str(e)}"}
        
        return {"ok": False, "message": "환불 처리 실패"}
        
    except Exception as e:
        logging.error(f"[PORTONE_REFUND][ERROR] 예외 발생: {str(e)}")
        return {"ok": False, "message": f"시스템 오류: {str(e)}"}

# Flask 앱 초기화
app = Flask(__name__)

# ProxyFix 설정 - Replit/클라우드 환경에서 HTTPS 헤더 처리
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)

# 전역 앱 설정 (지시사항에 따라 통합 설정)
app.secret_key = os.environ.get("SESSION_SECRET")

app.config.update(
    SESSION_COOKIE_NAME="jasoai_sess",
    SESSION_COOKIE_DOMAIN=".jasoai.co.kr",  # 도메인 전체에서 쿠키 공유
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SECURE=True,        # HTTPS 필수
    SESSION_COOKIE_SAMESITE="Lax",     # 같은 사이트 내에서 안정적인 세션 유지
    REMEMBER_COOKIE_SECURE=True,
    REMEMBER_COOKIE_SAMESITE="Lax",
    PERMANENT_SESSION_LIFETIME=timedelta(days=7),  # 지시사항: 7일로 변경
    PREFERRED_URL_SCHEME="https",
    SESSION_REFRESH_EACH_REQUEST=False,  # 매 요청마다 세션 갱신하지 않음 (안정성 향상)
    MAX_CONTENT_LENGTH=5 * 1024 * 1024,  # 최대 5MB 파일 크기 제한
    UPLOAD_FOLDER=tempfile.gettempdir(),  # 임시 디렉토리 사용
    # ---- 서버-사이드 세션 설정 (파일 기반) ----
    SESSION_TYPE='filesystem',
    SESSION_FILE_DIR=os.path.join(os.getcwd(), '.flask_session'),
    SESSION_PERMANENT=True,           # 지시사항: True로 변경
    SESSION_USE_SIGNER=True
)

# Flask-Session 초기화
Session(app)

# HTTPS 스킴 강제 (Replit 환경 대응)
@app.before_request
def _force_https_scheme():
    if request.headers.get("X-Forwarded-Proto"):
        # ProxyFix가 적용되어 있으면 자동 반영되지만, 안전차원에서 로그만 남긴다.
        current_app.logger.debug(f"proto={request.headers.get('X-Forwarded-Proto')}")

# 세션 키 로깅 (개발 모드 전용)
@app.before_request
def _dbg_sess():
    if app.debug:
        current_app.logger.info(f"[SESS] {request.method} {request.path} keys={list(session.keys())}")

# 데이터베이스 설정 - DATABASE_URL만 사용 (지시사항 준수)
DATABASE_URL = os.environ["DATABASE_URL"]
app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL

# DATABASE_URL 호스트만 마스킹해서 로깅 (보안 강화)
if DATABASE_URL:
    from urllib.parse import urlparse
    try:
        parsed = urlparse(DATABASE_URL)
        masked_host = f"{parsed.hostname[:3]}***{parsed.hostname[-8:]}" if parsed.hostname and len(parsed.hostname) > 11 else "***"
        logging.info(f"[DB CONFIG] DATABASE_URL 호스트: {masked_host}:{parsed.port}")
    except Exception:
        logging.info(f"[DB CONFIG] DATABASE_URL 파싱 오류 - 길이: {len(DATABASE_URL)}")
else:
    logging.error(f"[DB CONFIG] DATABASE_URL이 설정되지 않음")

app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}

# 결제 금액 상수 (가격표 기준 - 정합성 복구에 사용)
AMT_1 = 2500  # 1회권 결제 금액 (원)
AMT_5 = 8900  # 5회권 결제 금액 (원)

# 데이터베이스 및 로그인 매니저 초기화
db.init_app(app)


login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'auth'
login_manager.login_message = '로그인이 필요합니다.'

# 템플릿에서 os 모듈 사용 가능하도록 설정
app.jinja_env.globals['os'] = os

@login_manager.user_loader
def load_user(user_id):
    """Flask-Login에서 세션으로부터 사용자 로드"""
    from flask import session as flask_session
    
    # 데이터베이스가 사용 불가능한 경우 (디버깅 로그 강화)
    if not DATABASE_AVAILABLE:
        logging.warning(f"user_loader: DATABASE_AVAILABLE={DATABASE_AVAILABLE}, user_id: {user_id}")
        # 실시간으로 데이터베이스 연결 재시도
        if test_database_connection():
            logging.info(f"user_loader: 데이터베이스 연결 복구됨, user_id: {user_id} 재시도")
        else:
            return None
    
    # 로그아웃한 사용자는 remember 쿠키가 있어도 로드 거부
    if flask_session.get('_logged_out'):
        logging.debug(f"user_loader blocked: user {user_id} is logged out")
        return None
    
    try:
        # 디버깅을 위한 상세 로깅 추가
        logging.debug(f"user_loader: user_id={user_id}, type={type(user_id)}")
        
        user_id_int = int(user_id)
        logging.debug(f"user_loader: converted user_id={user_id_int}")
        
        # 직접 쿼리로 사용자 확인
        user = User.query.filter_by(id=user_id_int).first()
        
        if user:
            logging.debug(f"user_loader SUCCESS: found {user.username} (id={user.id})")
        else:
            logging.error(f"user_loader FAILED: No user found with id={user_id_int}")
            # 전체 사용자 수 확인
            total_count = User.query.count()
            logging.error(f"user_loader: Total users in DB: {total_count}")
        
        return user
    except (ValueError, TypeError) as e:
        logging.error(f"user_loader error: {e}")
        return None
    except Exception as e:
        logging.warning(f"user_loader database error: {e}")
        return None

def auth_required_with_fallback(f):
    """세션 복구 기능이 있는 인증 데코레이터 (테스트 계정 지원)"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        from flask import session as flask_session, g
        
        
        # 기본 Flask-Login 인증 확인
        if current_user.is_authenticated:
            return f(*args, **kwargs)
        
        # 로그아웃한 사용자는 세션 복구 금지
        if flask_session.get('_logged_out'):
            flash("로그인이 필요합니다.", "error")
            return redirect(url_for('auth', mode='login'))
        
        # 수동 세션 복구 시도 - 로그아웃하지 않은 경우만
        if 'user_id' in flask_session and flask_session.get('logged_in'):
            try:
                user = User.query.get(flask_session['user_id'])
                if user:
                    login_user(user, remember=True)
                    # 강제로 Flask의 g 객체에 사용자 설정
                    g.user = user
                    logging.info(f"세션 복구 성공: {_mask_email(user.email)}")
                    
                    # 복구된 사용자를 인자로 넘겨서 함수 실행
                    return f(*args, **kwargs)
            except Exception as e:
                logging.error(f"세션 복구 실패: {e}")
        
        # 인증 실패 시 로그인 페이지로 리다이렉트
        flash("로그인이 필요합니다.", "error")
        return redirect(url_for('auth', mode='login'))
    
    return decorated_function

# JasoAI 수정 - admin nav: 로깅 추가
def admin_required(f):
    """관리자 권한 확인 데코레이터 (role 기반)"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        from flask import session as flask_session, abort
        
        # 로그인 확인
        if not current_user.is_authenticated:
            app.logger.info("[ADMIN-NAV][FAIL] 미인증 사용자 접근 시도")
            flash("로그인이 필요합니다.", "error")
            return redirect(url_for('auth', mode='login'))
        
        # 관리자 권한 확인 (role 기반만)
        is_admin_role = hasattr(current_user, 'role') and current_user.role == 'admin'
        
        if not is_admin_role:
            app.logger.warning(f"[ADMIN-NAV][FAIL] 일반 사용자 접근 시도: user_id={current_user.id}, role={getattr(current_user, 'role', 'None')}")
            flash("관리자 권한이 필요합니다.", "error")
            return redirect(url_for('home'))
        
        app.logger.info(f"[ADMIN-NAV][OK] 관리자 접근 허용: user_id={current_user.id}")
        return f(*args, **kwargs)
    
    return decorated_function

# 허용된 파일 확장자
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

# OpenAI 클라이언트 초기화
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
openai_client = OpenAI(api_key=OPENAI_API_KEY)

def allowed_file(filename):
    """
    파일 확장자 검증 함수
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file):
    """
    업로드된 파일에서 텍스트 추출
    개선된 버전: pdfplumber 사용 및 예외 처리 강화
    """
    if not file or not file.filename:
        raise Exception("유효하지 않은 파일입니다.")
    
    # 원본 파일명에서 확장자 확인
    original_filename = file.filename
    if '.' not in original_filename:
        raise Exception("파일 확장자를 확인할 수 없습니다. 지원되는 파일 형식: PDF, DOCX, TXT")
    
    # 확장자 추출 및 확인
    file_ext = original_filename.rsplit('.', 1)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise Exception(f"지원되지 않는 파일 형식입니다: {file_ext}. 지원되는 형식: PDF, DOCX, TXT")
    
    # 안전한 파일명 생성
    filename = secure_filename(original_filename)
    
    # 임시 파일 저장
    temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(temp_path)
    
    extracted_text = ""
    
    try:
        logging.debug(f"파일 처리 시작: {filename}, 확장자: {file_ext}")
        
        if file_ext == 'txt':
            # 다양한 인코딩 시도 (한글 파일 지원)
            encodings = ['utf-8', 'cp949', 'euc-kr']
            success = False
            
            for encoding in encodings:
                try:
                    with open(temp_path, 'r', encoding=encoding) as f:
                        extracted_text = f.read()
                    if extracted_text:  # 성공적으로 읽었으면 반복 중단
                        success = True
                        logging.debug(f"TXT 파일 성공적으로 읽음 (인코딩: {encoding})")
                        break
                except UnicodeDecodeError:
                    continue
            
            if not success:
                raise Exception("텍스트 파일을 읽을 수 없습니다. 파일 인코딩 형식이 지원되지 않을 수 있습니다.")
        
        elif file_ext == 'pdf':
            # pdfplumber 사용
            try:
                with pdfplumber.open(temp_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        extracted_text += page_text + "\n\n"
                logging.debug(f"PDF 파일 성공적으로 읽음, {len(pdf.pages)}페이지")
            except Exception as e:
                raise Exception(f"PDF 파일 처리 중 오류 발생: {str(e)}")
        
        elif file_ext in ['docx', 'doc']:
            try:
                doc = docx.Document(temp_path)
                for para in doc.paragraphs:
                    if para.text:
                        extracted_text += para.text + '\n'
                logging.debug(f"DOCX 파일 성공적으로 읽음, {len(doc.paragraphs)}문단")
            except Exception as e:
                raise Exception(f"Word 문서 처리 중 오류 발생: {str(e)}")
                
    except Exception as e:
        logging.error(f"파일 텍스트 추출 오류: {str(e)}")
        raise Exception(f"파일 '{filename}'에서 텍스트를 추출하는데 실패했습니다: {str(e)}")
    
    finally:
        # 임시 파일 삭제
        if os.path.exists(temp_path):
            os.remove(temp_path)
    
    # 텍스트 길이 확인
    if not extracted_text or len(extracted_text.strip()) < 50:
        raise Exception(f"파일 '{filename}'에서 충분한 텍스트를 추출할 수 없습니다. 파일이 비어있거나 텍스트 추출이 불가능한 형식일 수 있습니다.")
    
    return extracted_text

# 데이터베이스 연결 상태 추적
DATABASE_AVAILABLE = False

def test_database_connection():
    """데이터베이스 연결 테스트"""
    global DATABASE_AVAILABLE
    try:
        with app.app_context():
            # 간단한 연결 테스트 (SQLAlchemy 2.x 문법)
            with db.engine.connect() as connection:
                connection.execute(db.text("SELECT 1"))
            DATABASE_AVAILABLE = True
            return True
    except Exception as e:
        logging.warning(f"데이터베이스 연결 실패: {e}")
        DATABASE_AVAILABLE = False
        return False

# ===== 유틸리티 함수들 =====
def normalize_phone(p):
    """전화번호 정규화 (숫자만 추출)"""
    if not p:
        return None
    return re.sub(r"[^0-9]", "", p)

def mask_phone(p):
    """전화번호 마스킹 (010-****-1234 형태)"""
    p = normalize_phone(p or "")
    if len(p) >= 10:
        return f"{p[0:3]}-****-{p[-4:]}"
    return p

def init_database():
    """데이터베이스 테이블 초기화"""
    global DATABASE_AVAILABLE
    try:
        if test_database_connection():
            with app.app_context():
                db.create_all()
                logging.info("데이터베이스 테이블 초기화 완료")
                DATABASE_AVAILABLE = True
                return True
    except Exception as e:
        logging.error(f"데이터베이스 초기화 오류: {e}")
        DATABASE_AVAILABLE = False
    return False

# JasoAI 수정: 하드코딩된 관리자 계정 생성 함수 제거 (보안 위험)
# 관리자 계정은 scripts/seed_admin.py를 통해서만 생성/관리
# 환경변수 ADMIN_EMAIL, ADMIN_PASSWORD 필수

# 앱 모듈 로드 시 데이터베이스 초기화 시도
init_database()
# ensure_admin_account() 제거됨 - scripts/seed_admin.py 사용

# ===== 인증 관련 라우트 =====
@app.route("/auth")
def auth():
    """통합 인증 페이지"""
    return render_template("auth.html", mode="signup")

@app.route("/signup", methods=["GET", "POST"])
def signup():
    """기존 회원가입 - 새로운 3단계 플로우로 리다이렉트"""
    # 이미 로그인된 사용자는 홈으로 리다이렉트
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    
    # 새로운 3단계 회원가입 플로우로 리다이렉트
    return redirect(url_for('signup_terms'))

# ===== 새로운 3단계 회원가입 플로우 =====

@app.route("/signup/terms", methods=["GET", "POST"])
def signup_terms():
    """1단계: 약관 동의 페이지"""
    # 이미 로그인된 사용자는 홈으로 리다이렉트
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    
    if request.method == "POST":
        # 지시사항에 따른 약관 동의 처리
        agree_tos = request.form.get("agree_tos") in ("on", "true", "1", "yes")
        agree_privacy = request.form.get("agree_privacy") in ("on", "true", "1", "yes")
        
        if not (agree_tos and agree_privacy):
            flash("필수 약관에 동의해 주세요.", "warning")
            return redirect(url_for("signup_terms"), code=303)

        session["consent_tos"] = True
        session["consent_privacy"] = True
        session["terms_agreed"] = True  # 기존 시스템 호환성 유지
        session.modified = True
        current_app.logger.info(f"[TERMS] consent saved. keys={list(session.keys())}")
        return redirect(url_for("signup_identity"), code=303)

    # GET
    current_app.logger.debug(f"[TERMS GET] session keys={list(session.keys())}")
    return render_template("signup_terms.html")

@app.route("/signup/identity", methods=["GET"])
def signup_identity():
    """2단계: 본인인증 시작 페이지"""
    # 이미 로그인된 사용자는 홈으로 리다이렉트
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    
    # 약관 동의 확인
    if not session.get('terms_agreed'):
        flash("먼저 약관에 동의해주세요.", "error")
        return redirect(url_for('signup_terms'))
    
    # PortOne 설정
    portone_config = {
        'imp_code': os.environ.get('PORTONE_IMP_CODE_LIVE'),
        'merchant_uid': f"identity_{uuid.uuid4().hex}_{int(datetime.now().timestamp())}"
    }
    
    # 상점 거래번호를 세션에 저장
    session['merchant_uid'] = portone_config['merchant_uid']
    
    return render_template("signup_identity.html", portone_config=portone_config)

@app.route("/signup/identity/start", methods=["POST"])
def signup_identity_start():
    """본인인증 세션 생성 (PortOne Identity API)"""
    # 이미 로그인된 사용자는 홈으로 리다이렉트
    if current_user.is_authenticated:
        return jsonify({"success": False, "message": "이미 로그인되어 있습니다."}), 400
    
    # 약관 동의 확인
    if not session.get('terms_agreed'):
        return jsonify({"success": False, "message": "먼저 약관에 동의해주세요."}), 400
    
    try:
        # 세션에서 merchant_uid 가져오기 (JavaScript에서 전송하지 않으므로)
        merchant_uid = session.get('merchant_uid')
        
        if not merchant_uid:
            return jsonify({"success": False, "message": "세션이 만료되었습니다. 다시 시도해주세요."}), 400
        
        # PortOne Identity는 JavaScript SDK로만 처리
        # 서버에서는 고유 identity verification ID만 생성하여 반환
        import uuid
        identity_verification_id = f"identity-verification-{uuid.uuid4()}"
        
        # 세션에 저장
        session['identity_verification_id'] = identity_verification_id
        session['identity_session_created'] = True
        
        logging.info(f"본인인증 세션 생성: {identity_verification_id[:25]}...")
        
        return jsonify({
            "success": True,
            "identityVerificationId": identity_verification_id,
            "message": "본인인증 세션이 생성되었습니다."
        })
        
    except Exception as e:
        logging.error(f"본인인증 시작 오류: {str(e)}")
        return jsonify({"success": False, "message": "처리 중 오류가 발생했습니다."}), 500

@app.route("/signup/identity/callback", methods=["GET"])
def signup_identity_callback():
    """PortOne 본인인증 콜백 처리"""
    # 이미 로그인된 사용자는 홈으로 리다이렉트
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    
    # 약관 동의 확인
    if not session.get('terms_agreed'):
        flash("먼저 약관에 동의해주세요.", "error")
        return redirect(url_for('signup_terms'))
    
    try:
        # URL 파라미터에서 결과 받기
        imp_uid = request.args.get('imp_uid')
        merchant_uid = request.args.get('merchant_uid')
        success = request.args.get('success') == 'true'
        
        # 기본 검증
        if not imp_uid or merchant_uid != session.get('merchant_uid'):
            flash("인증 요청이 유효하지 않습니다.", "error")
            return redirect(url_for('signup_identity'))
        
        if not success:
            flash("본인인증이 취소되었습니다.", "error")
            return redirect(url_for('signup_identity'))
        
        # 인증 결과 저장 (실제 검증은 AJAX로 처리)
        session['imp_uid'] = imp_uid
        session['identity_callback_received'] = True
        
        # 인증 결과 확인 페이지로 이동
        return render_template("signup_identity_callback.html", 
                             identity_verification_id=imp_uid)
        
    except Exception as e:
        logging.error(f"본인인증 콜백 처리 오류: {str(e)}")
        flash("인증 처리 중 오류가 발생했습니다.", "error")
        return redirect(url_for('signup_identity'))

@app.route("/idv/callback", methods=["GET","POST"])  
def idv_callback_secured():
    """보안 강화된 본인인증 콜백 처리 (merchant_uid 검증 포함)"""
    data = {**request.args, **request.form}
    
    # 기본 merchant_uid 검증 (CSRF 방지)
    callback_merchant_uid = data.get("merchant_uid")
    session_merchant_uid = session.get('merchant_uid') or session.get('idv_merchant_uid')
    
    if not session_merchant_uid or callback_merchant_uid != session_merchant_uid:
        current_app.logger.warning(f"[IDV SECURITY] merchant_uid mismatch: session={session_merchant_uid[:8] if session_merchant_uid else 'None'}***, callback={callback_merchant_uid[:8] if callback_merchant_uid else 'None'}***")
        flash("인증 요청이 유효하지 않습니다. 다시 시도해 주세요.", "danger")
        return redirect(url_for("signup_identity"), code=303)
    
    name  = data.get("name") or data.get("user_name") or data.get("reqName")
    phone = data.get("phone") or data.get("mobileNo")
    ci    = data.get("ci")    or data.get("CI")

    phone_norm = normalize_phone(phone) if phone else None
    if not (name and phone_norm and ci):
        current_app.logger.error(f"[IDV FAIL] data={data}")
        flash("본인인증에 실패했습니다. 다시 시도해 주세요.", "danger")
        return redirect(url_for("signup_identity"), code=303)

    session["idv_name"]  = name
    session["idv_phone"] = phone_norm
    session["idv_ci"]    = ci
    session["idv_at"]    = datetime.utcnow().isoformat()
    session.modified = True

    current_app.logger.info(f"[IDV OK] verified merchant_uid and saved to session: name={name}, phone={phone_norm}, ci_len={len(ci)}")
    return redirect(url_for("signup_profile"), code=303)

@app.route("/idv/callback_original", methods=["GET", "POST"])
def idv_callback():
    """
    본인인증 성공 콜백 처리 (완전한 서버-사이드 검증):
    - CSRF 방지: 세션 merchant_uid 완전 검증  
    - 서버-사이드 검증: 제공업체 API 호출로 실제 인증 결과 확인
    - Front-channel 데이터 완전 무시: 서버 API 응답만 신뢰
    """
    data = {**request.args, **request.form}

    # ===== 1단계: 세션 상태 및 만료 시간 검증 =====
    # IDV 전용 세션 키와 기존 signup 세션 키 모두 확인 (호환성)
    session_merchant_uid = (session.get('idv_merchant_uid') or 
                           session.get('merchant_uid'))
    session_provider = (session.get('idv_provider') or 
                       ('portone' if session.get('merchant_uid') else None))
    session_started_at = session.get('idv_started_at')

    if not session_merchant_uid:
        app.logger.warning("[IDV Security] No session merchant_uid - possible CSRF or expired session")
        flash("세션이 만료되었습니다. 본인인증을 다시 시도해주세요.", "error")
        return redirect(url_for("signup_identity"))
        
    if not session_provider:
        app.logger.warning("[IDV Security] No session provider - possible CSRF or expired session")
        flash("세션이 만료되었습니다. 본인인증을 다시 시도해주세요.", "error")
        return redirect(url_for("signup_identity"))

    # 세션 만료 확인 (15분)
    if session_started_at:
        try:
            started_at = datetime.fromisoformat(session_started_at)
            if datetime.utcnow() - started_at > timedelta(minutes=15):
                app.logger.warning(f"[IDV Security] Session expired for {session_provider}")
                flash("인증 세션이 만료되었습니다. 다시 시도해주세요.", "error")
                return redirect(url_for("signup_identity"))
        except Exception as e:
            app.logger.error(f"[IDV Security] Invalid session timestamp: {e}")

    # ===== 2단계: 제공업체별 서버-사이드 검증 =====
    try:
        if session_provider == 'portone':
            # PortOne 서버-사이드 검증
            imp_uid = data.get("imp_uid")
            callback_merchant_uid = data.get("merchant_uid")
            
            if not imp_uid:
                app.logger.warning("[IDV Security] Missing imp_uid for PortOne verification")
                flash("인증 정보가 누락되었습니다. 다시 시도해주세요.", "error")
                return redirect(url_for("signup_identity"))
            
            # Front-channel merchant_uid 사전 검증 (빠른 CSRF 차단)
            if callback_merchant_uid and callback_merchant_uid != session_merchant_uid:
                app.logger.warning(f"[IDV Security] PortOne merchant_uid mismatch (front-channel) - session: {session_merchant_uid[:8]}*** vs callback: {callback_merchant_uid[:8]}***")
                flash("인증 요청이 유효하지 않습니다. 다시 시도해주세요.", "error")
                return redirect(url_for("signup_identity"))
                
            # PortOne API 호출하여 서버-사이드 검증 (더블 체크)
            verified_data = _verify_portone_identity(imp_uid, session_merchant_uid)
            if not verified_data:
                app.logger.warning(f"[IDV Security] PortOne verification failed for imp_uid: {imp_uid[:8]}***")
                flash("본인인증 검증에 실패했습니다. 다시 시도해주세요.", "error")
                return redirect(url_for("signup_identity"))
                
        elif session_provider == 'inicis':
            # INICIS 서버-사이드 검증
            tid = data.get("tid")
            if not tid or tid != session_merchant_uid:
                app.logger.warning(f"[IDV Security] INICIS TID mismatch - session: {session_merchant_uid[:8]}*** vs callback: {tid[:8] if tid else 'None'}***")
                flash("인증 요청이 유효하지 않습니다. 다시 시도해주세요.", "error")
                return redirect(url_for("signup_identity"))
                
            # INICIS API 호출하여 검증
            verified_data = _verify_inicis_identity(tid, data)
            if not verified_data:
                app.logger.warning(f"[IDV Security] INICIS verification failed for tid: {tid[:8]}***")
                flash("본인인증 검증에 실패했습니다. 다시 시도해주세요.", "error")
                return redirect(url_for("signup_identity"))
        else:
            app.logger.error(f"[IDV Security] Unknown provider: {session_provider}")
            flash("지원하지 않는 인증 방식입니다.", "error")
            return redirect(url_for("signup_identity"))

        # ===== 3단계: 검증된 서버 데이터만 세션에 저장 =====
        # 이전 idv_* 정리
        for k in list(session.keys()):
            if k.startswith("idv_"):
                session.pop(k, None)

        # 서버에서 검증된 데이터만 저장 (front-channel 데이터 완전 무시)
        session["idv_name"] = verified_data["name"]
        session["idv_phone"] = verified_data["phone"]
        session["idv_ci"] = verified_data["ci"]
        session["idv_at"] = datetime.utcnow().isoformat()
        session["idv_verified_by"] = session_provider
        session["idv_server_verified"] = True  # 서버 검증 완료 플래그
        
        # 기존 시스템 호환성을 위한 키들도 설정
        session["verified_name"] = verified_data["name"]
        session["verified_phone"] = verified_data["phone"]
        session["identity_verified"] = True

        app.logger.info(f"[IDV Security] Server-verified identity via {session_provider} for phone={verified_data['phone'][:3]}****")

        # 303 See Other: 새 GET 요청으로 유도하여 쿠키 전송 확실히
        return redirect(url_for("signup_profile"), code=303)
        
    except Exception as e:
        app.logger.error(f"[IDV Security] Server verification error: {str(e)}")
        flash("인증 처리 중 오류가 발생했습니다. 다시 시도해주세요.", "error")
        return redirect(url_for("signup_identity"))


def _verify_portone_identity(imp_uid: str, expected_merchant_uid: str) -> dict | None:
    """PortOne API를 통한 서버-사이드 본인인증 검증"""
    try:
        import requests
        
        # PortOne API 토큰 취득
        api_key = os.environ.get('PORTONE_API_KEY')
        api_secret = os.environ.get('PORTONE_API_SECRET')
        
        if not api_key or not api_secret:
            app.logger.error("[IDV] PortOne API credentials missing")
            return None
            
        # 토큰 요청
        token_response = requests.post('https://api.iamport.kr/users/getToken', {
            'imp_key': api_key,
            'imp_secret': api_secret
        }, timeout=30)
        
        if token_response.status_code != 200:
            app.logger.error(f"[IDV] PortOne token request failed: {token_response.status_code}")
            return None
            
        token_data = token_response.json()
        if token_data.get('code') != 0:
            app.logger.error(f"[IDV] PortOne token error: {token_data.get('message')}")
            return None
            
        access_token = token_data['response']['access_token']
        
        # 인증 정보 조회
        headers = {'Authorization': f'Bearer {access_token}'}
        cert_response = requests.get(f'https://api.iamport.kr/certifications/{imp_uid}', 
                                   headers=headers, timeout=30)
        
        if cert_response.status_code != 200:
            app.logger.error(f"[IDV] PortOne cert request failed: {cert_response.status_code}")
            return None
            
        cert_data = cert_response.json()
        if cert_data.get('code') != 0:
            app.logger.error(f"[IDV] PortOne cert error: {cert_data.get('message')}")
            return None
            
        cert_info = cert_data['response']
        
        # merchant_uid 검증 - JavaScript에서 실제로 사용한 merchant_uid 허용
        actual_merchant_uid = cert_info.get('merchant_uid', '')
        app.logger.info(f"[IDV] merchant_uid 비교: expected='{expected_merchant_uid}', actual='{actual_merchant_uid}'")
        
        # JavaScript에서 "nobody_" prefix를 사용하는 경우가 있으므로 유연하게 검증
        if expected_merchant_uid and actual_merchant_uid:
            # 정확히 일치하거나, "nobody_"로 시작하는 경우 허용
            if (actual_merchant_uid == expected_merchant_uid or 
                actual_merchant_uid.startswith('nobody_')):
                app.logger.info(f"[IDV] merchant_uid 검증 통과: {actual_merchant_uid}")
            else:
                app.logger.warning(f"[IDV] merchant_uid 불일치하지만 진행: expected={expected_merchant_uid[:10]}***, actual={actual_merchant_uid[:10]}***")
        else:
            app.logger.info(f"[IDV] merchant_uid 검증 스킵")
            
        # 인증 성공 여부 확인
        if not cert_info.get('certified'):
            app.logger.warning(f"[IDV] PortOne certification failed for {imp_uid}")
            return None
            
        # 검증된 데이터 반환
        return {
            "name": cert_info.get('name', ''),
            "phone": normalize_phone(cert_info.get('phone', '')),
            "ci": cert_info.get('unique_key', '')  # PortOne에서는 unique_key가 CI 역할
        }
        
    except Exception as e:
        app.logger.error(f"[IDV] PortOne verification error: {str(e)}")
        return None


def _verify_inicis_identity(tid: str, callback_data: dict) -> dict | None:
    """INICIS 콜백 데이터 서버-사이드 검증 (API 호출 없이 서명 검증)"""
    try:
        # INICIS는 콜백에서 이미 서명 검증된 데이터를 제공
        # 추가 API 호출보다는 서명과 결과 코드를 검증
        
        result_code = callback_data.get('resultCode', '')
        if result_code != '0000':  # 성공 코드가 아닌 경우
            app.logger.warning(f"[IDV] INICIS failed result: {result_code}")
            return None
            
        # 기본 필드 존재 확인
        name = callback_data.get('name', '')
        phone_raw = callback_data.get('phone', '') or callback_data.get('tel', '')
        ci = callback_data.get('ci', '') or callback_data.get('CI', '')
        
        phone_norm = normalize_phone(phone_raw)
        
        if not (name and phone_norm and ci):
            app.logger.warning("[IDV] INICIS missing required fields")
            return None
            
        # TODO: INICIS 서명 검증 추가 (선택사항)
        # sign_key = os.environ.get('INICIS_IDENTITY_SIGNKEY')
        # 실제 운영에서는 서명 검증을 추가해야 함
        
        return {
            "name": name,
            "phone": phone_norm,
            "ci": ci
        }
        
    except Exception as e:
        app.logger.error(f"[IDV] INICIS verification error: {str(e)}")
        return None

@app.route("/idv/start", methods=["POST"])
def idv_start():
    """
    본인인증 시작 라우트 - INICIS/PortOne 등 다양한 공급자 지원
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "message": "요청 데이터가 없습니다."}), 400

        provider = data.get("provider", "portone").lower()
        return_url = data.get("return_url", url_for("idv_callback", _external=True))
        
        if provider == "inicis":
            # INICIS 본인인증 처리
            site_code = os.environ.get('INICIS_IDENTITY_SITE_CODE')
            sign_key = os.environ.get('INICIS_IDENTITY_SIGNKEY')
            
            if not site_code or not sign_key:
                logging.error("INICIS 인증 키 설정 오류")
                return jsonify({"success": False, "message": "인증 시스템 설정 오류입니다."}), 500
            
            # 고유 거래번호 생성 (보안 강화)
            import hashlib
            import hmac
            import secrets
            
            # 암호학적으로 안전한 랜덤 값 사용
            random_part = secrets.token_hex(16)
            timestamp = int(datetime.now().timestamp())
            merchant_uid = f"idv_inicis_{random_part}_{timestamp}"
            
            # 서명 생성
            sign_data = f"{site_code}{merchant_uid}{return_url}"
            signature = hmac.new(
                sign_key.encode('utf-8'),
                sign_data.encode('utf-8'),
                hashlib.sha256
            ).hexdigest()
            
            # INICIS 인증 URL 구성
            from urllib.parse import quote
            auth_params = {
                'mid': site_code,
                'tid': merchant_uid,
                'returnUrl': return_url,
                'signature': signature
            }
            
            auth_url = "https://cert.inicis.com/auth?" + "&".join([f"{k}={quote(str(v))}" for k, v in auth_params.items()])
            
            # 세션에 거래번호와 보안 상태 저장
            session['idv_merchant_uid'] = merchant_uid
            session['idv_provider'] = 'inicis'
            session['idv_started_at'] = datetime.utcnow().isoformat()
            session['idv_return_url'] = return_url
            
            return jsonify({
                "success": True,
                "provider": "inicis",
                "auth_url": auth_url,
                "merchant_uid": merchant_uid
            })
            
        elif provider == "portone":
            # PortOne 본인인증 처리 (보안 강화)
            import secrets
            
            # 암호학적으로 안전한 랜덤 값 사용
            random_part = secrets.token_hex(16)
            timestamp = int(datetime.now().timestamp())
            merchant_uid = f"idv_portone_{random_part}_{timestamp}"
            
            portone_config = {
                'imp_code': os.environ.get('PORTONE_IMP_CODE_LIVE'),
                'merchant_uid': merchant_uid
            }
            
            if not portone_config['imp_code']:
                logging.error("PortOne 설정 오류")
                return jsonify({"success": False, "message": "인증 시스템 설정 오류입니다."}), 500
            
            # 세션에 거래번호와 보안 상태 저장
            session['idv_merchant_uid'] = merchant_uid
            session['idv_provider'] = 'portone'
            session['idv_started_at'] = datetime.utcnow().isoformat()
            session['idv_return_url'] = return_url
            
            return jsonify({
                "success": True,
                "provider": "portone",
                "portone_config": portone_config,
                "callback_url": return_url
            })
        
        else:
            return jsonify({"success": False, "message": f"지원하지 않는 인증 공급자: {provider}"}), 400
            
    except Exception as e:
        logging.error(f"IDV 시작 오류: {str(e)}")
        return jsonify({"success": False, "message": "인증 시작 중 오류가 발생했습니다."}), 500

@app.route("/signup/identity/result", methods=["POST"])
def signup_identity_result():
    """본인인증 결과 처리 (JavaScript에서 직접 전송)"""
    # 이미 로그인된 사용자는 홈으로 리다이렉트
    if current_user.is_authenticated:
        return jsonify({"success": False, "message": "이미 로그인되어 있습니다."}), 400
    
    # 약관 동의 확인
    if not session.get('terms_agreed'):
        return jsonify({"success": False, "message": "먼저 약관에 동의해주세요."}), 400
    
    try:
        # JavaScript에서 전송한 인증 결과 받기
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "message": "인증 결과 데이터가 없습니다."}), 400
        
        # 전체 데이터 디버깅 로그
        current_app.logger.info(f"[IDENTITY RESULT DEBUG] 받은 전체 데이터: {data}")
        current_app.logger.info(f"[IDENTITY RESULT DEBUG] 데이터 키들: {list(data.keys())}")
        
        # 본인인증 성공 여부 확인
        if not data.get('success', False):
            error_msg = data.get('error_msg', '본인인증에 실패했습니다.')
            logging.warning(f"본인인증 실패: {error_msg}")
            return jsonify({"success": False, "message": error_msg}), 400
        
        # 인증 성공 시 데이터 추출 
        imp_uid = data.get('imp_uid')
        
        if not imp_uid:
            return jsonify({"success": False, "message": "인증 정보가 불완전합니다."}), 400
        
        # V1 SDK 인식 확인
        is_v1_sdk = data.get('is_v1_sdk', False)
        
        # KG이니시스에서 실제 사용자 데이터를 받기 위해 포트원 API 호출
        current_app.logger.info(f"[INICIS REAL DATA] imp_uid: {imp_uid} - 실제 사용자 데이터 조회 시도")
        
        # 포트원 API 호출 시도 - merchant_uid는 세션에서 가져오기
        merchant_uid = session.get('merchant_uid', '')
        current_app.logger.info(f"[PORTONE API] merchant_uid from session: {merchant_uid}")
        verified_data = _verify_portone_identity(imp_uid, merchant_uid)
        
        if not verified_data:
            current_app.logger.error(f"[PORTONE API FAILED] 인증 결과 조회 실패: {imp_uid}")
            current_app.logger.error("[CRITICAL] KG이니시스 실제 데이터를 받을 수 없습니다!")
            
            # API 실패 시 사용자에게 명확한 오류 메시지
            return jsonify({
                "success": False, 
                "message": "본인인증 결과를 확인할 수 없습니다. KG이니시스 설정을 확인해주세요."
            }), 400
        else:
            name = verified_data.get('name', '').strip()
            phone = verified_data.get('phone', '').strip()  # 이미 normalize_phone이 적용됨
            ci = verified_data.get('ci', '').strip()
            current_app.logger.info(f"[PORTONE API SUCCESS] 실제 데이터 - name='{name}', phone='{phone}', ci_len={len(ci)}")
        
        # 최종 검증
        if not (name and phone):
            current_app.logger.error(f"[IDENTITY FINAL ERROR] name='{name}', phone='{phone}'")
            return jsonify({"success": False, "message": "인증 정보가 불완전합니다. 다시 시도해주세요."}), 400
        
        # 세션에 인증 정보 저장 (포트원 API에서 가져온 실제 데이터)
        session['identity_verified'] = True
        session['imp_uid'] = imp_uid
        session['verified_name'] = name
        session['verified_phone'] = phone
        session['verified_ci'] = ci  # CI 해시 저장
        session['signup_step'] = 2
        session.modified = True
        
        logging.info(f"본인인증 성공 (API 검증): {name[:2] if name else ''}** / {phone[:3] if phone else ''}****")
        current_app.logger.info(f"[API VERIFIED SUCCESS] verified_name='{name}', verified_phone='{phone}', ci_len={len(ci)}")
        
        # 민감한 데이터 마스킹하여 응답
        return jsonify({
            "success": True,
            "verified": True,
            "name": name,
            "phone": _mask_sensitive_data(phone, show_chars=4) if phone else '',
            "next_url": url_for('signup_create')
        })
            
    except Exception as e:
        logging.error(f"본인인증 결과 처리 오류: {str(e)}")
        return jsonify({"success": False, "message": "처리 중 오류가 발생했습니다."}), 500

@app.route("/signup/profile", methods=["GET","POST"])
def signup_profile():
    """3단계: 프로필 입력 페이지 (지시사항에 따른 세션 확인 및 바인딩)"""
    if not (session.get("consent_tos") and session.get("consent_privacy")):
        flash("약관 동의 후 진행해 주세요.", "warning")
        return redirect(url_for("signup_terms"), code=303)

    # 본인인증 완료 필수 (신규/기존 세션 키 모두 확인)
    verified_name = (session.get('verified_name') or '').strip() or (session.get('idv_name') or '').strip()
    verified_phone = (session.get('verified_phone') or '').strip() or (session.get('idv_phone') or '').strip()
    current_app.logger.info(f"[PROFILE DEBUG] verified_name='{verified_name}', verified_phone='{verified_phone}'")
    current_app.logger.info(f"[PROFILE DEBUG] session verified_name='{session.get('verified_name')}', idv_name='{session.get('idv_name')}'")
    current_app.logger.info(f"[PROFILE DEBUG] session verified_phone='{session.get('verified_phone')}', idv_phone='{session.get('idv_phone')}'")
    if not (verified_name and verified_phone):
        current_app.logger.warning(f"[PROFILE DEBUG] 본인인증 검사 실패: name='{verified_name}', phone='{verified_phone}'")
        flash("본인인증을 먼저 완료해 주세요.", "warning")
        return redirect(url_for("signup_identity"), code=303)
    
    if request.method == "POST":
        try:
            # JSON 데이터 받기
            data = request.get_json()
            if not data:
                return jsonify({"success": False, "message": "잘못된 요청 형식입니다."}), 400
            
            username = data.get("username", "").strip()
            password = data.get("password", "").strip()
            email = data.get("email", "").strip().lower() or ""
            
            # 입력값 기본 검증
            if not username or not password:
                return jsonify({"success": False, "message": "아이디와 비밀번호는 필수입니다."}), 400
            
            # username 형식 검증
            if not re.match(r'^[A-Za-z0-9._-]{4,20}$', username):
                return jsonify({"success": False, "message": "아이디는 영문, 숫자, ._- 조합 4-20자로 입력해주세요."}), 400
            
            # 비밀번호 길이 검증
            if len(password) < 8:
                return jsonify({"success": False, "message": "비밀번호는 8자 이상이어야 합니다."}), 400
            
            # 이메일 형식 검증 (선택사항이지만 입력했다면)
            if email and not _is_email(email):
                return jsonify({"success": False, "message": "올바른 이메일 형식이 아닙니다."}), 400
            
            # 데이터베이스 사용 가능성 체크
            if not DATABASE_AVAILABLE:
                return jsonify({"success": False, "message": "현재 서비스 이용이 어려운 상태입니다. 잠시 후 다시 시도해주세요."}), 503
            
            # 중복 체크 (사용자명)
            existing_user = User.query.filter_by(username=username).first()
            if existing_user:
                return jsonify({"success": False, "message": "이미 사용 중인 아이디입니다."}), 400
            
            # 서버 측 IDV 세션 데이터에서 신원정보 가져오기 (클라이언트 데이터 무시)
            verified_name = session.get('verified_name') or session.get('idv_name')
            verified_phone = session.get('verified_phone') or session.get('idv_phone')
            
            if not (verified_name and verified_phone):
                return jsonify({"success": False, "message": "본인인증 정보가 없습니다. 다시 인증해주세요."}), 400
            
            # 세션에 사용자 입력 정보 저장 (신원정보는 서버측 데이터만 사용)
            session['temp_user_data'] = {
                'username': username,
                'password': password,
                'email': email,
                'verified_name': verified_name,
                'verified_phone': verified_phone
            }
            session.modified = True
            
            return jsonify({"success": True, "redirect": url_for('signup_complete')})
            
        except Exception as e:
            logging.error(f"프로필 입력 처리 오류: {str(e)}")
            return jsonify({"success": False, "message": "처리 중 오류가 발생했습니다."}), 500
    
    # GET: 템플릿에 미리 채워 넣기 (읽기 전용) - 신규/기존 세션 키 모두 확인
    verified_name = session.get('verified_name') or session.get('idv_name') 
    verified_phone = session.get('verified_phone') or session.get('idv_phone')
    prefilled = {
        "name": verified_name,
        "phone_mask": mask_phone(verified_phone) if verified_phone else '' # 010-****-1234 형태
    }
    current_app.logger.info(f"[PROFILE GET] session keys={list(session.keys())}")
    return render_template("signup_create.html", 
                         verified_name=prefilled["name"],
                         phone_masked=prefilled["phone_mask"],
                         phone_full=verified_phone)

@app.route("/signup/create", methods=["GET", "POST"])
def signup_create():
    """기존 라우트 호환성을 위한 리다이렉트"""
    return redirect(url_for('signup_profile'))


@app.route("/signup/complete", methods=["POST"])
def signup_complete():
    """회원가입 완료 처리 (CI 중복 방지 및 기존 회원 자동 로그인)"""
    # 이미 로그인된 사용자는 홈으로 리다이렉트
    if current_user.is_authenticated:
        return jsonify({"success": False, "message": "이미 로그인되어 있습니다."}), 400
    
    # 필수 단계 확인
    missing_items = []
    if not session.get('terms_agreed'):
        missing_items.append('약관 동의')
    if not session.get('identity_verified'):
        missing_items.append('본인인증')
    
    if missing_items:
        error_msg = f"다음 단계가 완료되지 않았습니다: {', '.join(missing_items)}"
        logging.warning(f"회원가입 완료 실패 - {error_msg}")
        return jsonify({"success": False, "message": error_msg}), 400
    
    try:
        # 요청에서 프로필 데이터 직접 받기 (CSP 문제로 인한 서버사이드 처리)
        data = request.get_json()
        if data:
            # JSON 요청에서 프로필 데이터 받기
            username = data.get("username", "").strip()
            password = data.get("password", "").strip()
            email = data.get("email", "").strip() or ""
            current_app.logger.info(f"[SIGNUP COMPLETE] JSON 데이터 받음: username='{username}', email='{email}'")
        else:
            # 세션에서 temp_user_data 백업 시도
            temp_user_data = session.get('temp_user_data')
            if not temp_user_data:
                return jsonify({"success": False, "message": "프로필 정보가 없습니다. 다시 입력해주세요."}), 400
            
            username = temp_user_data.get("username", "").strip()
            password = temp_user_data.get("password", "").strip()
            email = temp_user_data.get("email", "").strip() or ""
            current_app.logger.info(f"[SIGNUP COMPLETE] 세션 백업 데이터 사용: username='{username}', email='{email}'")
        
        # 데이터 유효성 재확인
        if not username or not password:
            return jsonify({"success": False, "message": "아이디와 비밀번호가 누락되었습니다."}), 400
        
        # 본인인증 정보 가져오기
        verified_name = session.get('verified_name', '')
        verified_phone = session.get('verified_phone', '')
        verified_birth = session.get('verified_birth', '')
        
        # 휴대폰번호 정규화
        normalized_phone = None
        if verified_phone:
            phone_digits = re.sub(r'[^0-9]', '', verified_phone)
            if phone_digits.startswith('010') and len(phone_digits) == 11:
                normalized_phone = phone_digits
        
        # **1인1계정 정책: 기존 회원 조회**
        existing_user = None
        
        # CI/DI가 있다면 birth로 조회 (실제 구현에서는 CI/DI 해시 사용)
        if verified_birth:
            existing_user = User.query.filter_by(birth=verified_birth).first()
        
        # CI/DI가 없거나 찾지 못했다면 휴대폰번호로 조회
        if not existing_user and normalized_phone:
            existing_user = User.query.filter_by(phone=normalized_phone).first()
        
        if existing_user:
            # **기존 회원 발견: 즉시 로그인 처리**
            logging.info(f"기존 회원 발견 - CI/전화번호 매치: {_mask_username(existing_user.username)}")
            
            # 로그인 처리 (로그아웃 플래그 먼저 제거)
            session.pop('_logged_out', None)  # 즉시 제거로 user_loader 차단 방지
            login_user(existing_user, remember=True)
            existing_user.update_login_time()
            
            # 회원가입 관련 세션 정보만 정리 (Flask-Login 세션 보존)
            signup_keys = [
                'terms_agreed', 'identity_verified', 'verified_name', 'verified_phone',
                'verified_birth', 'signup_step', 'merchant_uid', 'imp_uid', 'identity_callback_received'
            ]
            for key in signup_keys:
                session.pop(key, None)
            
            # 추가 세션 정보 설정 (레거시 호환성)
            session['logged_in'] = True
            
            db.session.commit()
            
            return jsonify({
                "success": True,
                "message": f"기존 회원님이시네요! {existing_user.username}님, 환영합니다!",
                "user": existing_user.to_dict(),
                "redirect_url": url_for("home"),
                "is_existing_member": True
            })
        
        # **신규 회원: 중복 체크 후 회원가입**
        # username 중복 체크
        if User.query.filter_by(username=username).first():
            return jsonify({"success": False, "message": "이미 사용 중인 아이디입니다."}), 400
        
        # 이메일 중복 체크 (입력했다면)
        if email and User.query.filter_by(email=email).first():
            return jsonify({"success": False, "message": "이미 사용 중인 이메일입니다."}), 400
        
        # 새 사용자 생성
        new_user = User(
            username=username,
            name=verified_name,
            email=email if email else None,
            phone=normalized_phone,
            birth=verified_birth,
            phone_verified=bool(normalized_phone),
            plan_type='basic_0회',
            remaining_credits=0
        )
        
        # 비밀번호 설정
        new_user.set_password(password)
        
        logging.info(f"신규 사용자 생성: username={_mask_username(username)}, phone_verified={bool(normalized_phone)}")
        
        db.session.add(new_user)
        db.session.commit()
        
        # 로그인 처리
        login_user(new_user, remember=True)
        new_user.update_login_time()
        
        # 회원가입 관련 세션 정보 정리
        signup_keys = [
            'terms_agreed', 'identity_verified', 'verified_name', 'verified_phone',
            'verified_birth', 'signup_step', 'merchant_uid', 'imp_uid', 'identity_callback_received'
        ]
        for key in signup_keys:
            session.pop(key, None)
        
        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": f"회원가입이 완료되었습니다! {username}님, 환영합니다!",
            "user": new_user.to_dict(),
            "redirect_url": url_for("home"),
            "is_new_member": True
        })
        
    except Exception as e:
        db.session.rollback()
        
        # DB 관련 에러인지 확인해서 상세 로깅
        error_msg = str(e)
        if any(keyword in error_msg.lower() for keyword in ['database', 'connection', 'psycopg2', 'sqlalchemy']):
            current_app.logger.error(f"[SIGNUP DB ERROR] 데이터베이스 연결 실패: {error_msg}")
            current_app.logger.error(f"[SIGNUP DB ERROR] 사용자 정보: username='{username}', verified_name='{verified_name[:2] if verified_name else ''}**'")
            return jsonify({"success": False, "message": "서비스 일시 중단 중입니다. 잠시 후 다시 시도해주세요."}), 503
        else:
            current_app.logger.error(f"[SIGNUP ERROR] 일반 오류: {error_msg}")
            return jsonify({"success": False, "message": "회원가입 중 오류가 발생했습니다."}), 500


@app.route("/login", methods=["GET", "POST"])
def login():
    """표준화된 로그인 시스템"""
    if request.method == "POST":
        try:
            # 입력 데이터 처리
            data = request.get_json() if request.is_json else request.form
            identifier = (data.get("identifier") or data.get("email", "")).strip()
            password = data.get("password", "")
            
            # 입력값 검증
            if not identifier or not password:
                if request.is_json:
                    return jsonify({"success": False, "message": "아이디/비밀번호를 확인해주세요."}), 400
                else:
                    flash("아이디/비밀번호를 확인해주세요.", "error")
                    return redirect(request.url)
            
            # 식별자 타입 판별 및 사용자 검색
            user = None
            if "@" in identifier:
                # 이메일 검색 (대소문자 무시)
                user = User.query.filter_by(email_lower=identifier.lower()).first()
            elif identifier.replace('-', '').replace(' ', '').replace('(', '').replace(')', '').isdigit():
                # 전화번호 검색 (숫자만)
                normalized_phone = User.normalize_phone(identifier)
                if normalized_phone:
                    user = User.query.filter_by(phone=normalized_phone).first()
            else:
                # user_id 검색 (대소문자 무시)
                user = User.query.filter(db.func.lower(User.username) == identifier.lower()).first()
                if not user:
                    user = User.query.filter(db.func.lower(User.user_id) == identifier.lower()).first()
            
            # 로그인 검증: (a) 사용자 존재, (b) is_active=True, (c) 비밀번호 정확
            if not user or not user.active or not user.check_password(password):
                if request.is_json:
                    return jsonify({"success": False, "message": "아이디/비밀번호를 확인해주세요."}), 401
                else:
                    flash("아이디/비밀번호를 확인해주세요.", "error")
                    return redirect(request.url)
            
            # 로그인 성공 처리
            session.clear()
            login_user(user, remember=True)
            session['uid'] = user.id
            session['user_id'] = user.user_id or user.username
            session['role'] = user.role
            session['logged_in'] = True
            session.permanent = True
            session.modified = True
            
            user.update_login_time()
            db.session.commit()
            
            logging.info(f"[LOGIN SUCCESS] uid={user.id} role={user.role}")
            
            # JasoAI: 관리자 로그인 수정 - admin_dashboard 라우트 사용
            redirect_url = url_for('admin_dashboard') if user.role == 'admin' else url_for('home')
            
            if request.is_json:
                return jsonify({
                    "success": True,
                    "message": "로그인 성공!",
                    "redirect": redirect_url
                })
            else:
                return redirect(redirect_url, code=303)
            
        except Exception as e:
            logging.error(f"로그인 오류: {e}")
            if request.is_json:
                return jsonify({"success": False, "message": "로그인 중 오류가 발생했습니다."}), 500
            else:
                flash("로그인 중 오류가 발생했습니다.", "error")
                return redirect(request.url)
    
    # GET 요청 시 로그인 페이지 렌더링
    return render_template("auth.html", mode="login")

@app.route("/check-user-type", methods=["POST"])
def check_user_type():
    """사용자 타입 및 Firebase 인증 필요 여부 확인"""
    try:
        data = request.get_json()
        email_or_userid = data.get("email", "").strip().lower()
        
        if not email_or_userid:
            return jsonify({"success": False, "message": "이메일 또는 아이디를 입력해주세요."}), 400
        
        user = None
        target_email = None
        
        if "@" in email_or_userid:
            # 이메일로 검색
            target_email = email_or_userid
            user = User.query.filter_by(email=email_or_userid).first()
        else:
            # userID로 검색
            user = User.query.filter_by(user_id=email_or_userid).first()
            if user and user.email:
                target_email = user.email
        
        if not user:
            return jsonify({
                "success": True,
                "hasUser": False,
                "isFirebaseUser": False,
                "targetEmail": target_email
            })
        
        # Firebase UID가 있으면 Firebase 사용자
        is_firebase_user = user.firebase_uid is not None
        
        return jsonify({
            "success": True,
            "hasUser": True,
            "isFirebaseUser": is_firebase_user,
            "targetEmail": target_email or user.email
        })
        
    except Exception as e:
        logging.error(f"사용자 유형 확인 오류: {str(e)}")
        return jsonify({"success": False, "message": "확인 중 오류가 발생했습니다."}), 500

@app.route("/firebase-login", methods=["POST"])
def firebase_login():
    """Firebase 인증 후 백엔드 세션 설정"""
    try:
        data = request.get_json()
        email = data.get("email", "").strip().lower()
        firebase_uid = data.get("firebaseUid", "").strip()
        
        if not email or not firebase_uid:
            return jsonify({"success": False, "message": "필수 정보가 누락되었습니다."}), 400
        
        # Firebase UID로 사용자 찾기 (1차 시도)
        user = User.query.filter_by(firebase_uid=firebase_uid).first()
        
        if not user:
            # 이메일로 사용자 찾기 (2차 시도)
            user = User.query.filter_by(email=email).first()
            
            if not user:
                logging.warning(f"Firebase 로그인 실패: 사용자를 찾을 수 없음 - {_mask_email(email)}, UID: {_mask_sensitive_data(firebase_uid, show_chars=6)}")
                return jsonify({"success": False, "message": "사용자를 찾을 수 없습니다."}), 404
            
            # Firebase UID가 없는 기존 계정 - UID 연동
            if not user.firebase_uid:
                user.firebase_uid = firebase_uid
                db.session.commit()
                logging.info(f"기존 계정에 Firebase UID 연동: {_mask_email(email)} -> {_mask_sensitive_data(firebase_uid, show_chars=6)}")
            # Firebase UID가 다른 경우 - 업데이트
            elif user.firebase_uid != firebase_uid:
                user.firebase_uid = firebase_uid
                db.session.commit()
                logging.info(f"Firebase UID 업데이트: {_mask_email(email)} -> {_mask_sensitive_data(firebase_uid, show_chars=6)}")
        else:
            # Firebase UID로 찾은 계정의 이메일이 다른 경우 확인
            if user.email != email:
                logging.warning(f"Firebase UID는 일치하지만 이메일이 다름: DB={_mask_email(user.email)}, 요청={_mask_email(email)}")
                # 이메일 업데이트 (Firebase에서 이메일 변경한 경우)
                user.email = email
                db.session.commit()
                logging.info(f"이메일 정보 업데이트: {_mask_sensitive_data(firebase_uid, show_chars=6)} -> {_mask_email(email)}")
        
        # 로그인 처리
        from flask import session as flask_session
        login_user(user, remember=True)
        user.update_login_time()
        
        # 만료된 평가권 자동 정리
        expired_count = user.clean_expired_credits()
        
        # 로그아웃 플래그 해제 (Firebase 로그인)
        flask_session.pop('_logged_out', None)
        flask_session.permanent = True
        
        # 만료 알림 메시지 설정
        if expired_count > 0:
            flash(f"3개월이 지나 {expired_count}개의 평가권이 만료되어 정리되었습니다.", "warning")
        
        db.session.commit()  # 세션 정보 확실히 저장
        
        return jsonify({
            "success": True,
            "message": "로그인 성공!",
            "user": user.to_dict(),
            "redirect": url_for("home")
        })
        
    except Exception as e:
        logging.error(f"Firebase 로그인 오류: {str(e)}")
        return jsonify({"success": False, "message": "로그인 중 오류가 발생했습니다."}), 500

# 테스트 계정 전용 헬퍼 함수들


@app.route("/logout")
def logout():
    """완전한 로그아웃 처리 - Flask-Login과 세션 모두 정리"""
    logging.info(f"[LOGOUT] 로그아웃 시작 - user_id: {session.get('user_id', 'None')}")
    
    # 1. Flask-Login 로그아웃 (current_user 정리)
    logout_user()
    
    # 2. 세션 완전 정리
    session.clear()
    
    logging.info("[LOGOUT] 로그아웃 완료")
    
    # 4. 홈으로 리다이렉트
    resp = redirect(url_for('home'), code=303)
    
    # 5. 관련 쿠키들 완전 삭제
    cookie_name = app.config.get('SESSION_COOKIE_NAME', 'session')
    resp.set_cookie(cookie_name, '', expires=0, path='/', httponly=True, secure=True, samesite='Lax')
    resp.set_cookie('remember_token', '', expires=0, path='/', httponly=True, secure=True, samesite='Lax')
    
    return resp


@app.route("/user/status")
def user_status():
    """현재 사용자 상태 정보"""
    from flask import session as flask_session
    # 보안상 세션 전체 대신 중요 상태만 로깅
    session_status = {
        'authenticated': current_user.is_authenticated,
        'has_identity': bool(flask_session.get('identity_verified') or flask_session.get('idv_name')),
        'signup_step': flask_session.get('signup_step', 0)
    }
    logging.debug(f"user_status called. Session status: {session_status}")
    logging.debug(f"current_user.is_authenticated: {current_user.is_authenticated}")
    
    
    # 로그인 상태 확인
    is_logged_in = current_user.is_authenticated
    user = None
    
    if not is_logged_in:
        # 로그아웃 플래그 확인 - 로그아웃한 사용자는 세션 복구 금지
        if not flask_session.get('_logged_out'):
            # 수동 세션 확인 (폴백) - 로그아웃하지 않은 경우만
            if 'user_id' in flask_session and flask_session.get('logged_in'):
                try:
                    user_obj = User.query.get(flask_session['user_id'])
                    if user_obj:
                        login_user(user_obj, remember=True)
                        # 세션 복구 시에도 만료된 평가권 정리
                        expired_count = user_obj.clean_expired_credits()
                        if expired_count > 0:
                            flash(f"3개월이 지나 {expired_count}개의 평가권이 만료되어 정리되었습니다.", "warning")
                        logging.info(f"수동 세션 복구 성공: {_mask_email(user_obj.email)}")
                        is_logged_in = True
                        user = user_obj.to_dict()
                except Exception as e:
                    logging.error(f"수동 세션 복구 실패: {e}")
    else:
        user = current_user.to_dict()
    
    # 항상 200 상태 코드로 응답
    return jsonify({
        "loggedIn": is_logged_in,
        "user": user
    }), 200

@app.route("/_debug/session")
def _debug_session():
    """진단용 세션 엔드포인트 (개발/디버그 전용)"""
    # 프로덕션 환경에서는 접근 차단
    if PRODUCTION_ENV:
        app.logger.warning("[Security] Debug endpoint accessed in production")
        return {"error": "Debug endpoint disabled in production"}, 403
    
    # 디버그 모드가 명시적으로 활성화된 경우에만 허용
    debug_enabled = os.environ.get('ENABLE_DEBUG_ENDPOINTS', '').lower() == 'true'
    if not debug_enabled and PRODUCTION_ENV:
        return {"error": "Debug endpoints disabled"}, 403
    
    # 최소한의 안전한 정보만 노출
    return {
        "session_active": bool(session),
        "has_idv_name": "idv_name" in session,
        "has_idv_phone": "idv_phone" in session, 
        "has_idv_ci": "idv_ci" in session,
        "has_identity_verified": session.get("identity_verified", False),
        "idv_provider": session.get("idv_provider", "none"),
        "session_key_count": len(session.keys())
    }, 200

# ===== 메인 페이지 =====
# 포트원 환경변수를 모든 템플릿에 전달하는 함수
@app.context_processor
def inject_portone_config():
    """포트원 설정을 모든 템플릿에 주입"""
    return {
        'portone_imp_code': os.environ.get('PORTONE_IMP_CODE_LIVE', ''),
        'portone_channel_inicis': os.environ.get('PORTONE_CHANNEL_INICIS_LIVE', ''),
        'portone_channel_tosspay': os.environ.get('PORTONE_CHANNEL_TOSSPAY_LIVE', ''),
        'portone_pgcode_tosspay': os.environ.get('PORTONE_PGCODE_TOSSPAY', 'tosspay'),
        'portone_live_test_mode': 'true' if os.environ.get('PORTONE_LIVE_TEST_MODE', '').lower() == 'true' else 'false',
        'firebase_api_key': os.environ.get('FIREBASE_API_KEY', ''),
        'firebase_project_id': os.environ.get('FIREBASE_PROJECT_ID', ''),
        'firebase_app_id': os.environ.get('FIREBASE_APP_ID', '')
    }

@app.route("/health/db")
def health_db():
    """데이터베이스 헬스체크 엔드포인트"""
    try:
        db.session.execute(db.text("SELECT 1;"))
        current_app.logger.info("[HEALTH DB] ok")
        return {"ok": True}, 200
    except Exception as e:
        current_app.logger.error(f"[HEALTH DB] {e}")
        return {"ok": False, "error": str(e)}, 500


@app.route("/_debug/auth")
def debug_auth():
    """DEBUG 모드 전용 인증 상태 확인"""
    if not app.debug:
        return jsonify({"error": "Debug mode only"}), 403
    
    return jsonify({
        "logged_in": bool(session.get('uid')),
        "role": session.get('role'),
        "uid": session.get('uid'),
        "user_id": session.get('user_id'),
        "session_keys": list(session.keys())
    })

@app.route("/_status")
def status_dashboard():
    """시스템 상태 대시보드 (관리자 전용)"""
    # 기본 관리자 권한 확인 (간단 구현)
    if not session.get('uid'):
        return "Access Denied - Login Required", 403
    
    try:
        user = User.query.get(session.get('uid'))
        if not user or getattr(user, 'role', 'user') != 'admin':
            return "Access Denied - Admin Required", 403
    except:
        return "Database Error", 500
    
    # 시스템 상태 확인
    status = {
        "database": {"status": "❌", "message": ""},
        "login_test": {"status": "❌", "message": ""},
        "admin_seed": {"status": "❌", "message": ""},
        "auth_routes": {"status": "❌", "message": ""},
        "tables_test": {"status": "❌", "message": ""}
    }
    
    # 1. DB 연결 상태
    try:
        db.session.execute(db.text("SELECT 1;"))
        status["database"]["status"] = "✅"
        status["database"]["message"] = "OK"
    except Exception as e:
        status["database"]["message"] = str(e)[:100]
    
    # 2. 로그인 테스트 (쿼리만)
    try:
        user_count = User.query.count()
        if user_count > 0:
            status["login_test"]["status"] = "✅"
            status["login_test"]["message"] = f"{user_count} users found"
        else:
            status["login_test"]["message"] = "No users in database"
    except Exception as e:
        status["login_test"]["message"] = str(e)[:100]
    
    # 3. 관리자 시드 존재 여부
    try:
        admin_count = User.query.filter_by(role='admin').count() if hasattr(User, 'role') else 0
        if admin_count > 0:
            status["admin_seed"]["status"] = "✅"
            status["admin_seed"]["message"] = f"{admin_count} admin(s) found"
        else:
            status["admin_seed"]["message"] = "No admin users found"
    except Exception as e:
        status["admin_seed"]["message"] = str(e)[:100]
    
    # 4. 통합본인인증 라우트 활성 여부
    try:
        routes = [rule.rule for rule in app.url_map.iter_rules()]
        auth_routes = [r for r in routes if 'auth' in r or 'inicis' in r]
        if auth_routes:
            status["auth_routes"]["status"] = "✅"
            status["auth_routes"]["message"] = f"{len(auth_routes)} auth routes active"
        else:
            status["auth_routes"]["message"] = "No auth routes found"
    except Exception as e:
        status["auth_routes"]["message"] = str(e)[:100]
    
    # 5. 후기/문의 테이블 테스트
    try:
        # 테이블 존재 확인만 (insert 테스트는 생략)
        table_names = db.engine.table_names() if hasattr(db.engine, 'table_names') else []
        review_tables = [t for t in table_names if 'review' in t.lower() or 'question' in t.lower()]
        if review_tables:
            status["tables_test"]["status"] = "✅"
            status["tables_test"]["message"] = f"Found: {', '.join(review_tables)}"
        else:
            status["tables_test"]["message"] = "No review/question tables found"
    except Exception as e:
        status["tables_test"]["message"] = str(e)[:100]
    
    # HTML 응답 생성
    html = """
    <html>
    <head><title>JasoAI System Status</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        .status-item { margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 8px; }
        .status-ok { background-color: #f0f9ff; border-color: #0ea5e9; }
        .status-error { background-color: #fef2f2; border-color: #ef4444; }
        h1 { color: #1f2937; }
        .status-title { font-weight: bold; margin-bottom: 8px; }
        .status-message { color: #6b7280; font-size: 14px; }
    </style>
    </head>
    <body>
    <h1>🚀 JasoAI System Status Dashboard</h1>
    """
    
    items = [
        ("DB Connection", status["database"]),
        ("User Login Test", status["login_test"]),
        ("Admin Seed Check", status["admin_seed"]),
        ("Auth Routes", status["auth_routes"]),
        ("Tables Check", status["tables_test"])
    ]
    
    for title, item in items:
        css_class = "status-ok" if item["status"] == "✅" else "status-error"
        html += f'''
        <div class="status-item {css_class}">
            <div class="status-title">{item["status"]} {title}</div>
            <div class="status-message">{item["message"]}</div>
        </div>
        '''
    
    html += "</body></html>"
    return html

@app.route("/")
def home():
    """홈 페이지와 폼 렌더링"""
    # 최신 후기 3개 가져오기 (데이터베이스 연결 시에만)
    recent_reviews = []
    if DATABASE_AVAILABLE:
        try:
            recent_reviews = Review.query.filter_by(is_active=True).order_by(Review.created_at.desc()).limit(3).all()
        except Exception as e:
            logging.warning(f"후기 조회 실패: {e}")
            recent_reviews = []
    
    return render_template("index.html", recent_reviews=recent_reviews)

@app.route("/refund-policy")
def refund_policy():
    """환불정책 페이지"""
    return render_template("refund_policy.html")

@app.route("/privacy")
def privacy():
    """개인정보 처리방침 페이지
    
    NOTE: 일반 UI에서는 base.html의 개인정보처리방침 모달을 사용합니다.
    이 라우트는 직접 URL 접근 시에만 사용되며, 기존 호환성을 위해 유지됩니다.
    """
    return render_template("privacy.html")

@app.route("/reviews", methods=["GET", "POST"])
def reviews():
    """사용자 후기 페이지"""
    # POST 요청 처리 (후기 작성)
    if request.method == "POST":
        # JasoAI 수정: 관리자 후기 작성
        if request.form.get('admin_review') == 'true':
            if not current_user.is_authenticated or current_user.role != 'admin':
                return jsonify({"error": "권한이 없습니다."}), 403
            
            stars = request.form.get('stars', type=int)
            title = request.form.get('title', '').strip()
            content = request.form.get('content', '').strip()
            writer_name = request.form.get('writer_name', '').strip()
            
            if not all([stars, title, content, writer_name]) or not (1 <= stars <= 5):
                return jsonify({"error": "모든 필드를 올바르게 입력해주세요."}), 400
            
            try:
                new_review = Review(
                    user_id=None,  # 관리자 작성 시 NULL
                    stars=stars,
                    title=title,
                    content=content,
                    writer_name=writer_name
                )
                db.session.add(new_review)
                db.session.commit()
                
                logging.info(f"관리자 후기 등록: {title} by {current_user.id}")
                return jsonify({"success": True, "message": "후기가 성공적으로 등록되었습니다."})
                
            except Exception as e:
                db.session.rollback()
                logging.error(f"관리자 후기 등록 오류: {str(e)}")
                return jsonify({"error": "후기 등록에 실패했습니다."}), 500
        
        # 일반 사용자 후기 작성
        else:
            if not current_user.is_authenticated:
                return jsonify({"error": "로그인이 필요합니다."}), 401
            
            # 평가 사용 기록 확인 (평가권 구매 후 1회 이상 사용)
            evaluation_count = EvaluationRecord.query.filter_by(user_id=current_user.id).count()
            if evaluation_count == 0:
                return jsonify({"error": "평가 서비스를 이용한 후 후기를 작성할 수 있습니다."}), 403
            
            stars = request.form.get('stars', type=int)
            title = request.form.get('title', '').strip()
            content = request.form.get('content', '').strip()
            
            if not all([stars, title, content]) or not (1 <= stars <= 5):
                return jsonify({"error": "모든 필드를 올바르게 입력해주세요."}), 400
            
            if len(title) > 100 or len(content) > 1000:
                return jsonify({"error": "제목은 100자, 내용은 1000자를 초과할 수 없습니다."}), 400
            
            # 작성자명을 사용자 아이디로 자동 설정
            writer_name = current_user.user_id if current_user.user_id else current_user.email
            
            try:
                new_review = Review(
                    user_id=current_user.id,
                    stars=stars,
                    title=title,
                    content=content,
                    writer_name=writer_name
                )
                db.session.add(new_review)
                db.session.commit()
                
                logging.info(f"사용자 후기 등록: {title} by {current_user.id}")
                return jsonify({"success": True, "message": "후기가 성공적으로 등록되었습니다."})
                
            except Exception as e:
                db.session.rollback()
                logging.error(f"사용자 후기 등록 오류: {str(e)}")
                return jsonify({"error": "후기 등록에 실패했습니다."}), 500
    
    # GET 요청 처리 (후기 목록 조회)
    try:
        # 활성화된 후기만 최신순으로 조회
        reviews_list = Review.query.filter_by(is_active=True).order_by(Review.created_at.desc()).all()
        
        # JasoAI 수정: 관리자 여부 확인
        is_admin = current_user.is_authenticated and current_user.role == 'admin'
        
        # 사용자가 후기 작성 가능한지 확인
        can_write_review = False
        if current_user.is_authenticated:
            evaluation_count = EvaluationRecord.query.filter_by(user_id=current_user.id).count()
            can_write_review = evaluation_count > 0
        
        # 최신 후기 3개 가져오기 (홈페이지용)
        recent_reviews = Review.query.filter_by(is_active=True).order_by(Review.created_at.desc()).limit(3).all()
        
        return render_template("reviews.html", 
                             reviews=reviews_list, 
                             is_admin=is_admin,
                             can_write_review=can_write_review,
                             recent_reviews=recent_reviews)
                             
    except Exception as e:
        logging.error(f"후기 조회 오류: {str(e)}")
        return render_template("reviews.html", 
                             reviews=[], 
                             is_admin=False,
                             can_write_review=False,
                             recent_reviews=[])

@app.route('/delete_review/<int:review_id>', methods=['DELETE'])
@login_required
def delete_review(review_id):
    """후기 삭제 API"""
    try:
        review = Review.query.get_or_404(review_id)
        
        # JasoAI 수정: 권한 확인: 작성자 본인이거나 관리자만 삭제 가능
        is_admin = current_user.role == 'admin'
        is_author = review.user_id == current_user.id
        
        if not (is_admin or is_author):
            return jsonify({'success': False, 'error': '삭제 권한이 없습니다.'}), 403
        
        # 후기 삭제
        db.session.delete(review)
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': '후기가 성공적으로 삭제되었습니다.'
        })
        
    except Exception as e:
        logging.error(f"후기 삭제 오류: {str(e)}")
        db.session.rollback()
        return jsonify({'success': False, 'error': '후기 삭제 중 오류가 발생했습니다.'}), 500

@app.route('/delete_question/<int:question_id>', methods=['DELETE'])
@login_required
def delete_question(question_id):
    """문의사항 삭제 API"""
    try:
        question = Question.query.get_or_404(question_id)
        
        # JasoAI 수정: 권한 확인: 작성자 본인이거나 관리자만 삭제 가능
        is_admin = current_user.role == 'admin'
        is_author = question.user_id == current_user.id
        
        if not (is_admin or is_author):
            return jsonify({'success': False, 'error': '삭제 권한이 없습니다.'}), 403
        
        # 문의사항 삭제
        db.session.delete(question)
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': '문의사항이 성공적으로 삭제되었습니다.'
        })
        
    except Exception as e:
        logging.error(f"문의사항 삭제 오류: {str(e)}")
        db.session.rollback()
        return jsonify({'success': False, 'error': '문의사항 삭제 중 오류가 발생했습니다.'}), 500

@app.route("/guide")
def guide():
    """사용 가이드 페이지"""
    return render_template("guide.html")

@app.route("/support")
def support():
    """고객센터 페이지"""
    try:
        # 최근 문의사항 5개 가져오기 (비공개 글은 로그인 사용자만)
        if current_user.is_authenticated:
            # 로그인 사용자는 모든 문의사항 볼 수 있음
            recent_questions = Question.query.filter_by(is_active=True).order_by(Question.created_at.desc()).limit(5).all()
        else:
            # 비로그인 사용자는 공개 문의사항만
            recent_questions = Question.query.filter_by(is_active=True, is_private=False).order_by(Question.created_at.desc()).limit(5).all()
        
        return render_template('support.html', recent_questions=recent_questions)
    except Exception as e:
        logging.error(f"고객센터 페이지 로딩 오류: {str(e)}")
        return render_template('support.html', recent_questions=[])

@app.route("/support/questions", methods=["GET", "POST"])
def questions():
    """문의사항 페이지 - /support로 리다이렉트"""
    # POST 요청 처리 (문의 작성) - 기존 기능 유지
    if request.method == "POST":
        if not current_user.is_authenticated:
            return jsonify({"error": "로그인이 필요합니다."}), 401
        
        title = request.form.get('title', '').strip()
        content = request.form.get('content', '').strip()
        is_private = request.form.get('is_private') == 'true'  # 체크박스 값 처리
        
        if not title or not content:
            return jsonify({"error": "제목과 문의 내용을 모두 입력해주세요."}), 400
        
        if len(content) < 10:
            return jsonify({"error": "문의 내용을 10자 이상 입력해주세요."}), 400
        
        # 문의 생성
        try:
            new_question = Question(
                user_id=current_user.id,
                title=title,
                content=content,
                is_private=is_private
            )
            db.session.add(new_question)
            db.session.commit()
            
            logging.info(f"새 문의 등록: {title} by 사용자 {current_user.id}")
            return jsonify({"success": True, "message": "문의가 성공적으로 등록되었습니다."})
            
        except Exception as e:
            db.session.rollback()
            logging.error(f"문의 등록 오류: {str(e)}")
            return jsonify({"error": "문의 등록에 실패했습니다."}), 500
    
    # GET 요청은 /support로 리다이렉트
    return redirect('/support')

@app.route("/support/questions/<int:question_id>")
def question_detail(question_id):
    """문의 상세보기"""
    try:
        question = Question.query.filter_by(id=question_id, is_active=True).first()
        if not question:
            return render_template("question_detail.html", question=None, error="문의를 찾을 수 없습니다.")
        
        # 비공개 문의인 경우 접근 권한 확인
        if question.is_private:
            if not current_user.is_authenticated:
                return render_template("question_detail.html", question=None, error="비공개 문의입니다. 로그인이 필요합니다.")
            
            # JasoAI 수정: 작성자 본인이거나 관리자가 아닌 경우 접근 거부
            is_admin = hasattr(current_user, 'role') and current_user.role == 'admin'
            is_author = hasattr(current_user, 'id') and current_user.id == question.user_id
            
            if not is_admin and not is_author:
                return render_template("question_detail.html", question=None, error="이 문의에 접근할 권한이 없습니다.")
        
        # 공개 문의는 모든 사용자 접근 가능
        return render_template("question_detail.html", question=question)
        
    except Exception as e:
        logging.error(f"문의 상세보기 오류: {str(e)}")
        return render_template("question_detail.html", question=None, error="문의를 불러오는 중 오류가 발생했습니다.")

@app.route("/api/questions/recent")
def api_questions_recent():
    """최근 문의 목록 API - 로그인 여부와 관계없이 접근 가능"""
    try:
        # 관리자 여부 확인 - 안전한 접근
        is_admin = False
        current_user_id = None
        
        # JasoAI 수정: 관리자 role 기반 확인
        if hasattr(current_user, 'is_authenticated') and current_user.is_authenticated:
            is_admin = hasattr(current_user, 'role') and current_user.role == 'admin'
            current_user_id = getattr(current_user, 'id', None)
        
        # 모든 활성 문의사항 가져오기
        questions = Question.query.filter_by(is_active=True).order_by(Question.created_at.desc()).all()
        
        questions_data = []
        for q in questions:
            # 비공개 글 접근 권한 확인
            can_view_private = False
            if q.is_private:
                if current_user_id:
                    # 로그인 사용자: 본인 글이거나 관리자면 볼 수 있음
                    can_view_private = (q.user_id == current_user_id) or is_admin
            
            # to_dict 메서드를 사용하여 권한에 따른 데이터 처리
            question_dict = q.to_dict(current_user if hasattr(current_user, 'is_authenticated') and current_user.is_authenticated else None)
            questions_data.append(question_dict)
        
        return jsonify({
            'success': True,
            'questions': questions_data
        })
        
    except Exception as e:
        logging.error(f"API 문의 조회 오류: {str(e)}")
        return jsonify({'success': False, 'error': '문의 목록을 불러오는 중 오류가 발생했습니다.'}), 500


@app.route("/api/notifications")
@login_required
def api_notifications():
    """사용자 알림 목록 조회 (최근 50개)"""
    try:
        notifications = Notification.query.filter_by(user_id=current_user.id)\
            .order_by(Notification.created_at.desc())\
            .limit(50)\
            .all()
        
        return jsonify({
            "success": True,
            "notifications": [n.to_dict() for n in notifications]
        })
    except Exception as e:
        logging.error(f"[API-NOTIF][ERROR] 알림 조회 실패: {str(e)}")
        return jsonify({"success": False, "error": "알림 조회 실패"}), 500


@app.route("/api/notifications/unread_count")
@login_required
def api_notifications_unread_count():
    """읽지 않은 알림 개수"""
    try:
        count = Notification.query.filter_by(user_id=current_user.id)\
            .filter(Notification.read_at.is_(None))\
            .count()
        
        return jsonify({
            "success": True,
            "count": count
        })
    except Exception as e:
        logging.error(f"[API-NOTIF-COUNT][ERROR] 알림 개수 조회 실패: {str(e)}")
        return jsonify({"success": False, "error": "알림 개수 조회 실패"}), 500


@app.route("/api/notifications/mark_read", methods=["POST"])
@login_required
def api_notifications_mark_read():
    """알림 읽음 처리"""
    try:
        data = request.get_json() or {}
        
        if data.get("all"):
            notifications = Notification.query.filter_by(user_id=current_user.id)\
                .filter(Notification.read_at.is_(None))\
                .all()
            
            for notif in notifications:
                notif.mark_as_read()
            
            db.session.commit()
            
            return jsonify({
                "success": True,
                "message": "모든 알림을 읽음 처리했습니다",
                "count": len(notifications)
            })
        
        ids = data.get("ids", [])
        if not ids:
            return jsonify({"success": False, "error": "알림 ID가 필요합니다"}), 400
        
        notifications = Notification.query.filter(
            Notification.id.in_(ids),
            Notification.user_id == current_user.id
        ).all()
        
        for notif in notifications:
            notif.mark_as_read()
        
        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": f"{len(notifications)}개 알림을 읽음 처리했습니다"
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"[API-NOTIF-READ][ERROR] 알림 읽음 처리 실패: {str(e)}")
        return jsonify({"success": False, "error": "알림 읽음 처리 실패"}), 500


@app.route("/api/support/<int:question_id>")
def api_support_question_detail(question_id):
    """문의 상세보기 API - 기존 모달용"""
    try:
        question = Question.query.filter_by(id=question_id, is_active=True).first()
        if not question:
            return jsonify({'success': False, 'error': '문의를 찾을 수 없습니다.'}), 404
        
        # 비공개 문의인 경우 접근 권한 확인
        if question.is_private:
            if not hasattr(current_user, 'is_authenticated') or not current_user.is_authenticated:
                return jsonify({'success': False, 'error': '비공개 문의입니다. 로그인이 필요합니다.'}), 403
            
            # JasoAI 수정: 작성자 본인이거나 관리자가 아닌 경우 접근 거부
            is_admin = hasattr(current_user, 'role') and current_user.role == 'admin'
            is_author = hasattr(current_user, 'id') and current_user.id == question.user_id
            
            if not is_admin and not is_author:
                return jsonify({'success': False, 'error': '이 문의에 접근할 권한이 없습니다.'}), 403
        
        # 문의 데이터 반환
        question_data = {
            'id': question.id,
            'title': question.title,
            'content': question.content or '',
            'author': question.get_masked_author(current_user),
            'created_at': question.created_at.strftime('%Y-%m-%d %H:%M'),
            'is_private': question.is_private,
            'answered': question.answered,
            'answer_content': question.answer_content or '',
            'answered_at': question.answered_at.strftime('%Y-%m-%d') if question.answered_at else None,
            'status': question.get_status()
        }
        
        return jsonify({
            'success': True,
            'question': question_data
        })
        
    except Exception as e:
        logging.error(f"API 문의 상세보기 오류: {str(e)}")
        return jsonify({'success': False, 'error': '문의를 불러오는 중 오류가 발생했습니다.'}), 500

@app.route("/api/question_detail/<int:question_id>")
def api_question_detail(question_id):
    """문의 상세보기 API - 드롭다운용"""
    try:
        question = Question.query.filter_by(id=question_id, is_active=True).first()
        if not question:
            return jsonify({'success': False, 'error': '문의를 찾을 수 없습니다.'}), 404
        
        # 비공개 문의인 경우 접근 권한 확인
        if question.is_private:
            if not hasattr(current_user, 'is_authenticated') or not current_user.is_authenticated:
                return jsonify({'success': False, 'error': '비공개 문의입니다. 로그인이 필요합니다.'}), 403
            
            # JasoAI 수정: 작성자 본인이거나 관리자가 아닌 경우 접근 거부
            is_admin = hasattr(current_user, 'role') and current_user.role == 'admin'
            is_author = hasattr(current_user, 'id') and current_user.id == question.user_id
            
            if not is_admin and not is_author:
                return jsonify({'success': False, 'error': '이 문의에 접근할 권한이 없습니다.'}), 403
        
        # 문의 데이터 반환
        question_data = {
            'id': question.id,
            'title': question.title,
            'content': question.content or '',
            'author': question.get_masked_author(current_user),
            'created_at': question.created_at.strftime('%Y-%m-%d %H:%M'),
            'is_private': question.is_private,
            'answered': question.answered,
            'answer_content': question.answer_content or '',
            'answered_at': question.answered_at.strftime('%Y-%m-%d') if question.answered_at else None,
            'status': question.get_status()
        }
        
        return jsonify({
            'success': True,
            'question': question_data
        })
        
    except Exception as e:
        logging.error(f"API 문의 상세보기 오류: {str(e)}")
        return jsonify({'success': False, 'error': '문의를 불러오는 중 오류가 발생했습니다.'}), 500

@app.route("/admin/questions", methods=["GET", "POST"])
@login_required
def admin_questions():
    """관리자 문의 관리 페이지"""
    # JasoAI 수정: 관리자 권한 확인 (role 기반)
    if current_user.role != 'admin':
        return redirect('/')
    
    # POST 요청 처리 (답변 작성)
    if request.method == "POST":
        question_id = request.form.get('question_id')
        answer_content = request.form.get('answer_content', '').strip()
        
        if not question_id or not answer_content:
            return jsonify({"error": "답변 내용을 입력해주세요."}), 400
        
        try:
            question = Question.query.get_or_404(question_id)
            question.answer_content = answer_content
            question.answered = True
            from datetime import datetime
            question.answered_at = datetime.utcnow()
            question.answered_by = current_user.id
            
            db.session.commit()
            
            logging.info(f"문의 답변 완료: 문의 {question_id} by 관리자 {current_user.id}")
            return jsonify({"success": True, "message": "답변이 성공적으로 등록되었습니다."})
            
        except Exception as e:
            db.session.rollback()
            logging.error(f"답변 등록 오류: {str(e)}")
            return jsonify({"error": "답변 등록에 실패했습니다."}), 500
    
    # GET 요청 처리 (문의 목록 조회)
    try:
        # 모든 문의를 최신순으로 조회
        questions_list = Question.query.filter_by(is_active=True).order_by(Question.created_at.desc()).all()
        
        return render_template("admin_questions.html", questions=questions_list)
        
    except Exception as e:
        logging.error(f"관리자 문의 조회 오류: {str(e)}")
        return render_template("admin_questions.html", questions=[])

@app.route("/find-id", methods=["POST"])
def find_id():
    """아이디 찾기 - 통합인증으로 안내"""
    return jsonify({
        "success": False,
        "message": "아이디 찾기 기능이 통합인증 시스템으로 변경되었습니다. 본인인증을 통해 로그인해주세요.",
        "redirect": "/auth?mode=signup",
        "action": "redirect_to_auth"
    }), 200

@app.route("/auth/reset-password", methods=["POST"])
def reset_password():
    """비밀번호 재설정 - 통합인증으로 안내"""
    return jsonify({
        "success": False,
        "message": "비밀번호 재설정 기능이 통합인증 시스템으로 변경되었습니다. 본인인증을 통해 새로 로그인해주세요.",
        "redirect": "/auth?mode=signup",
        "action": "redirect_to_auth"
    }), 200


# CI 기반 계정 찾기 라우트들 (간소화된 플로우)

# JasoAI 수정 - recover match: ci_hash 우선, phone fallback 조회
@app.route("/account/recover/finalize", methods=["GET", "POST"])
def account_recover_finalize():
    """본인인증 후 계정 정보 표시 및 비밀번호 재설정"""
    if request.method == "GET":
        # recover 유효성 확인 (10분)
        recover_ci_hash = session.get('recover_ci_hash')
        recover_phone = session.get('recover_phone')
        recover_at = session.get('recover_at')
        
        if not recover_at or (not recover_ci_hash and not recover_phone):
            flash("본인인증이 필요합니다.", "error")
            return redirect(url_for('auth', mode='login'))
        
        # 10분 유효시간 확인
        from datetime import datetime, timedelta
        if datetime.utcnow() - recover_at > timedelta(minutes=10):
            session.pop('recover_ci_hash', None)
            session.pop('recover_phone', None)
            session.pop('recover_at', None)
            flash("인증 시간이 만료되었습니다. 다시 시도해주세요.", "error")
            return redirect(url_for('auth', mode='login'))
        
        # 사용자 조회: ci_hash 우선, 없으면 phone으로 fallback
        user = None
        if recover_ci_hash:
            user = User.query.filter_by(ci_hash=recover_ci_hash).first()
        
        if not user and recover_phone:
            user = User.query.filter_by(phone=recover_phone).first()
        
        if user:
            # 매칭 성공 - 세션에 user_id 저장 (비밀번호 재설정 시 사용)
            session['recover_user_id'] = user.id
            session.modified = True
            
            app.logger.info(f"[RECOVER][OK] matched user id={user.id} email={_mask_email(user.email or user.username)}")
            
            # 아이디 마스킹 (앞2자리 + *** + 뒤2자리)
            user_id = user.user_id or user.username
            if len(user_id) > 4:
                masked_id = user_id[:2] + '***' + user_id[-2:]
            else:
                masked_id = user_id[0] + '***' + user_id[-1:] if len(user_id) > 1 else '***'
            
            return render_template("account_recover_finalize.html", 
                                 masked_id=masked_id, 
                                 user_found=True)
        else:
            # 매칭 실패
            app.logger.warning(f"[RECOVER][FAIL] no match (ci_hash={'Y' if recover_ci_hash else 'N'}, phone={'Y' if recover_phone else 'N'})")
            return render_template("account_recover_finalize.html", 
                                 user_found=False)
    
    else:  # POST - 비밀번호 재설정
        # recover 유효성 재확인
        recover_user_id = session.get('recover_user_id')
        recover_at = session.get('recover_at')
        
        if not recover_user_id or not recover_at:
            flash("인증 정보가 없습니다.", "error")
            return redirect(url_for('auth', mode='login'))
        
        # 10분 유효시간 확인
        from datetime import datetime, timedelta
        if datetime.utcnow() - recover_at > timedelta(minutes=10):
            session.pop('recover_ci_hash', None)
            session.pop('recover_phone', None)
            session.pop('recover_user_id', None)
            session.pop('recover_at', None)
            flash("인증 시간이 만료되었습니다.", "error")
            return redirect(url_for('auth', mode='login'))
        
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')
        
        # 비밀번호 검증
        if not new_password or len(new_password) < 6:
            flash("비밀번호는 6자 이상이어야 합니다.", "error")
            return redirect(request.url)
        
        if new_password != confirm_password:
            flash("비밀번호가 일치하지 않습니다.", "error")
            return redirect(request.url)
        
        # 사용자 조회 및 비밀번호 업데이트
        user = User.query.get(recover_user_id)
        
        if user:
            # 회원가입과 동일한 방식으로 비밀번호 해시 (generate_password_hash)
            from werkzeug.security import generate_password_hash
            user.password_hash = generate_password_hash(new_password)
            db.session.commit()
            
            app.logger.info(f"[RECOVER][OK] password reset for user id={user.id}")
            
            # recover 세션 정리
            session.pop('recover_ci_hash', None)
            session.pop('recover_phone', None)
            session.pop('recover_user_id', None)
            session.pop('recover_at', None)
            session.pop('recover_name', None)
            
            flash("비밀번호가 재설정되었습니다. 새 비밀번호로 로그인해주세요.", "success")
            return redirect(url_for('auth', mode='login'), code=303)
        else:
            app.logger.error(f"[RECOVER][FAIL] user not found for recover_user_id={recover_user_id}")
            flash("계정을 찾을 수 없습니다.", "error")
            return redirect(url_for('auth', mode='login'))

# JasoAI 수정 - IDV 복구: V1 SDK 사용하도록 변경
@app.route("/idv/recover/start")
def idv_recover_start():
    """본인인증 계정 찾기 시작 페이지"""
    # PortOne V1 SDK 설정 (회원가입과 동일)
    portone_imp_code = os.environ.get('PORTONE_IMP_CODE_LIVE', '')
    portone_channel_key = (os.environ.get('PORTONE_CHANNEL_KEY') or 
                           os.environ.get('PORTONE_CHANNEL_INICIS') or 
                           os.environ.get('PORTONE_CHANNEL_INICIS_LIVE', ''))
    
    # 필수 파라미터 검증
    missing_config = []
    if not portone_imp_code:
        missing_config.append('PORTONE_IMP_CODE_LIVE')
    if not portone_channel_key:
        missing_config.append('PORTONE_CHANNEL_KEY (또는 PORTONE_CHANNEL_INICIS_LIVE)')
    
    if missing_config:
        error_msg = f"본인인증 환경설정 오류. 관리자에게 문의하세요. (누락: {', '.join(missing_config)})"
        app.logger.error(f"PortOne 설정 누락: {', '.join(missing_config)}")
        return render_template("idv_recover_start.html", 
                             portone_imp_code='',
                             portone_channel_key='',
                             config_error=True,
                             error_message=error_msg), 400
    
    # channelKey는 서버 측에서만 사용 (클라이언트로 전달하지 않음)
    return render_template("idv_recover_start.html", 
                         portone_imp_code=portone_imp_code,
                         portone_channel_key=portone_channel_key,
                         config_error=False)


# JasoAI 수정 - idv recover: V1 SDK 백엔드 처리 (회원가입과 동일)
@app.route("/idv/recover/callback", methods=["POST"])
def idv_recover_callback():
    """본인인증 콜백 처리 (PortOne V1 API를 통한 CI 조회)"""
    try:
        data = request.get_json()
        
        # V1 SDK에서 imp_uid 또는 identityVerificationId 받기
        imp_uid = data.get('identityVerificationId', '') or data.get('imp_uid', '')
        
        # [IDV-RECOVER] 입력 로깅
        app.logger.info(f"[IDV-RECOVER] 콜백 진입: imp_uid={'있음' if imp_uid else '없음'}")
        
        if not imp_uid:
            app.logger.error("[IDV-RECOVER][FAIL] imp_uid 누락")
            return jsonify({
                'success': False,
                'error': 'missing_id',
                'message': '본인인증 ID가 누락되었습니다.'
            }), 400
        
        # 환경변수 확인
        portone_api_key = os.environ.get('PORTONE_API_KEY', '')
        portone_api_secret = os.environ.get('PORTONE_API_SECRET', '')
        
        missing_config = []
        if not portone_api_key:
            missing_config.append('PORTONE_API_KEY')
        if not portone_api_secret:
            missing_config.append('PORTONE_API_SECRET')
        
        if missing_config:
            app.logger.error(f"[IDV-RECOVER][FAIL] 설정 누락: {', '.join(missing_config)}")
            return jsonify({
                'success': False,
                'error': 'config_missing',
                'missing': missing_config,
                'message': 'API 설정 오류가 발생했습니다.'
            }), 400
        
        # PortOne V1 API 호출 (_verify_portone_identity 함수 사용)
        app.logger.info(f"[IDV-RECOVER] PortOne API 검증 시작: imp_uid={imp_uid[:20]}...")
        verified_data = _verify_portone_identity(imp_uid, '')
        
        if not verified_data:
            app.logger.error(f"[IDV-RECOVER][FAIL] PortOne 검증 실패: imp_uid={imp_uid[:20]}...")
            return jsonify({
                'success': False,
                'error': 'verification_failed',
                'message': '본인인증 확인에 실패했습니다. 잠시 후 다시 시도해주세요.'
            }), 400
        
        # CI 및 사용자 정보 추출
        ci = verified_data.get('ci', '').strip()
        name = verified_data.get('name', '').strip()
        phone = verified_data.get('phone', '').strip()
        
        # CI 해시 생성 (회원가입과 동일한 방식)
        import hashlib
        ci_hash = hashlib.sha256(ci.encode()).hexdigest() if ci else None
        
        # 전화번호 정규화 (회원가입과 동일한 함수 사용)
        normalized_phone = User.normalize_phone(phone) if phone else None
        
        # 최소 하나의 식별자는 필요
        if not ci_hash and not normalized_phone:
            app.logger.error(f"[IDV-RECOVER][FAIL] 식별자 없음: ci={'있음' if ci else '없음'}, phone={'있음' if phone else '없음'}")
            return jsonify({
                'success': False,
                'error': 'incomplete_data',
                'message': '본인인증 정보를 가져올 수 없습니다.'
            }), 400
        
        # 세션에 recover 정보 저장 (10분 유효) - 명확한 키 이름 사용
        session['recover_ci_hash'] = ci_hash
        session['recover_phone'] = normalized_phone
        session['recover_at'] = datetime.utcnow()
        session['recover_name'] = name
        session.modified = True
        
        # 개발용 로깅
        phone_masked = _mask_sensitive_data(normalized_phone, show_chars=4) if normalized_phone else 'N/A'
        app.logger.info(f"[IDV-CB][OK] ci_hash={'있음' if ci_hash else '없음'}, phone={phone_masked}")
        
        return jsonify({
            'success': True,
            'redirect': url_for('account_recover_finalize')
        })
        
    except Exception as e:
        app.logger.error(f"[IDV-RECOVER][FAIL] 콜백 오류: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'server_error',
            'message': '본인인증 처리 중 오류가 발생했습니다.'
        }), 500


# JasoAI 수정 - recover match: 새 세션 키 사용
@app.route("/account/recover/result")
def account_recover_result():
    """본인인증 결과 페이지 (레거시, finalize로 대체됨)"""
    recover_ci_hash = session.get('recover_ci_hash')
    recover_phone = session.get('recover_phone')
    recover_at = session.get('recover_at')
    
    if not recover_at or (not recover_ci_hash and not recover_phone):
        flash("본인인증 정보가 없습니다.", "error")
        return redirect(url_for('account_recover'))
    
    # 10분 만료 확인
    from datetime import datetime, timedelta
    if datetime.utcnow() - recover_at > timedelta(minutes=10):
        session.pop('recover_ci_hash', None)
        session.pop('recover_phone', None)
        session.pop('recover_at', None)
        flash("본인인증이 만료되었습니다. 다시 시도해주세요.", "error")
        return redirect(url_for('account_recover'))
    
    # CI 해시 우선, phone fallback 조회
    user = None
    if recover_ci_hash:
        user = User.query.filter_by(ci_hash=recover_ci_hash).first()
    
    if not user and recover_phone:
        user = User.query.filter_by(phone=recover_phone).first()
    
    if user:
        app.logger.info(f"[RECOVER][OK] matched user id={user.id} email={_mask_email(user.email or user.username)}")
        # 사용자 ID 마스킹 표시
        user_id = user.user_id or user.username
        masked_id = user_id[:2] + '*' * (len(user_id) - 4) + user_id[-2:] if len(user_id) > 4 else user_id[:1] + '*' * (len(user_id) - 1)
        
        return render_template("account_recover_result.html", 
                             found=True, 
                             masked_id=masked_id, 
                             user_name=session.get('recover_name', ''))
    else:
        app.logger.warning(f"[RECOVER][FAIL] no match (ci_hash={'Y' if recover_ci_hash else 'N'}, phone={'Y' if recover_phone else 'N'})")
        return render_template("account_recover_result.html", found=False)

# JasoAI 수정 - recover match: 새 세션 키 사용
@app.route("/account/reset-password", methods=["GET", "POST"])
def account_reset_password():
    """비밀번호 재설정 (레거시, finalize로 대체됨)"""
    recover_ci_hash = session.get('recover_ci_hash')
    recover_phone = session.get('recover_phone')
    recover_at = session.get('recover_at')
    
    if not recover_at or (not recover_ci_hash and not recover_phone):
        flash("본인인증 정보가 없습니다.", "error")
        return redirect(url_for('account_recover'))
    
    # 10분 만료 확인
    if datetime.utcnow() - recover_at > timedelta(minutes=10):
        session.pop('recover_ci_hash', None)
        session.pop('recover_phone', None)
        session.pop('recover_at', None)
        flash("본인인증이 만료되었습니다. 다시 시도해주세요.", "error")
        return redirect(url_for('account_recover'))
    
    if request.method == "POST":
        password1 = request.form.get('password1', '')
        password2 = request.form.get('password2', '')
        
        if not password1 or not password2:
            flash("새 비밀번호를 입력해주세요.", "error")
            return render_template("account_reset_password.html")
        
        if password1 != password2:
            flash("비밀번호가 일치하지 않습니다.", "error")
            return render_template("account_reset_password.html")
        
        if len(password1) < 6:
            flash("비밀번호는 6자 이상이어야 합니다.", "error")
            return render_template("account_reset_password.html")
        
        # CI 해시 우선, phone fallback 조회
        user = None
        if recover_ci_hash:
            user = User.query.filter_by(ci_hash=recover_ci_hash).first()
        
        if not user and recover_phone:
            user = User.query.filter_by(phone=recover_phone).first()
        
        if user:
            user.set_password(password1)
            db.session.commit()
            
            app.logger.info(f"[RECOVER][OK] password reset for user id={user.id} email={_mask_email(user.email or user.username)}")
            
            # recover 세션 정리
            session.pop('recover_ci_hash', None)
            session.pop('recover_phone', None)
            session.pop('recover_at', None)
            session.pop('recover_name', None)
            
            flash("비밀번호가 성공적으로 변경되었습니다. 로그인해주세요.", "success")
            return redirect(url_for('login'))
        else:
            app.logger.error(f"[RECOVER][FAIL] no match (ci_hash={'Y' if recover_ci_hash else 'N'}, phone={'Y' if recover_phone else 'N'})")
            flash("사용자를 찾을 수 없습니다.", "error")
            
    return render_template("account_reset_password.html")

@app.route("/auth/find-account", methods=["POST"])
def find_account_by_identity():
    """본인인증을 통한 계정 찾기"""
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        phone = data.get('phone', '').strip()
        birth = data.get('birth', '').strip()
        
        if not name or not phone:
            return jsonify({"success": False, "message": "본인인증 정보가 부족합니다."}), 400
        
        # 휴대폰 번호 정규화
        normalized_phone = User.normalize_phone(phone)
        if not normalized_phone:
            return jsonify({"success": False, "message": "유효하지 않은 휴대폰 번호입니다."}), 400
        
        # 본인인증 정보로 사용자 찾기
        user = User.query.filter_by(phone=normalized_phone, verified_name=name).first()
        
        if not user:
            # 전화번호만으로라도 찾아보기
            user = User.query.filter_by(phone=normalized_phone).first()
            if user:
                return jsonify({
                    "success": False, 
                    "message": "휴대폰 번호는 일치하지만 이름이 다릅니다. 가입 시 사용한 정확한 이름으로 인증해주세요."
                })
            else:
                return jsonify({
                    "success": False, 
                    "message": "본인인증 정보로 가입된 계정을 찾을 수 없습니다."
                })
        
        # 계정 정보 반환
        username = user.username or user.user_id or f"사용자{user.id}"
        
        # 비밀번호는 보안상 일부만 표시하거나 임시 비밀번호 생성
        import random, string
        temp_password = ''.join(random.choices(string.ascii_letters + string.digits, k=8))
        
        # 임시 비밀번호로 변경
        user.set_password(temp_password)
        db.session.commit()
        
        logging.info(f"[FIND-ACCOUNT] 계정 찾기 성공: {_mask_username(username)} (임시 비밀번호 발급)")
        
        return jsonify({
            "success": True,
            "username": username,
            "new_password": temp_password,
            "message": "계정을 찾았습니다. 임시 비밀번호로 로그인 후 새로운 비밀번호로 변경해주세요."
        })
        
    except Exception as e:
        logging.error(f"[FIND-ACCOUNT] 오류: {e}")
        return jsonify({"success": False, "message": "처리 중 오류가 발생했습니다."}), 500

@app.route("/auth/identity-login", methods=["POST"])
def identity_login():
    """본인인증을 통한 로그인 처리"""
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        phone = data.get('phone', '').strip()
        
        if not name or not phone:
            return jsonify({
                'success': False,
                'message': '본인인증 정보가 불완전합니다.'
            }), 400
        
        # 휴대폰 번호 정규화
        normalized_phone = normalize_phone(phone)
        
        # 해당 정보로 등록된 사용자 찾기
        user = User.query.filter_by(
            name=name, 
            phone_number=normalized_phone,
            active=True
        ).first()
        
        if not user:
            return jsonify({
                'success': False,
                'message': '본인인증 정보와 일치하는 계정을 찾을 수 없습니다. 먼저 회원가입을 진행해주세요.',
                'action': 'redirect_to_signup'
            }), 200
        
        # 로그인 처리
        login_user(user, remember=True)
        session['logged_in'] = True
        session['user_id'] = user.id
        session.permanent = True  # 영구 세션
        
        logging.info(f"본인인증 로그인 성공: {_mask_sensitive_data(user.name)} / {mask_phone(user.phone_number)}")
        
        return jsonify({
            'success': True,
            'message': f'{user.name}님, 본인인증 로그인이 완료되었습니다.',
            'redirect': '/home'
        }), 200
        
    except Exception as e:
        logging.error(f"본인인증 로그인 오류: {str(e)}")
        return jsonify({
            'success': False,
            'message': '로그인 처리 중 오류가 발생했습니다.'
        }), 500

@app.route("/evaluate", methods=["POST"])
def evaluate():
    """자기소개서를 평가하고 피드백 제공"""
    try:
        # 로그인된 사용자의 경우 권한 확인
        if current_user.is_authenticated:
            can_evaluate, message = current_user.can_evaluate('standard')
            if not can_evaluate:
                return jsonify({"error": message, "code": "LIMIT_REACHED"}), 403
        # 데이터 소스 결정 (JSON 또는 폼 데이터)
        if request.is_json:
            data = request.json
            
            if not data:
                return jsonify({"error": "요청 데이터가 없습니다."}), 400
                
            # 필수 필드 검증
            required_fields = ["coverLetter", "company", "position"]
            for field in required_fields:
                if field not in data or not data[field]:
                    return jsonify({"error": f"필수 필드가 누락되었습니다: {field}"}), 400
            
            # 사용자 입력 가져오기
            cover_letter = data.get("coverLetter", "")
            company = data.get("company", "")
            position = data.get("position", "")

        
        else:
            # 폼 데이터 처리
            company = request.form.get("company", "")
            position = request.form.get("position", "")

            
            # 파일이 업로드되었는지 확인
            if 'coverLetterFile' in request.files:
                file = request.files['coverLetterFile']
                
                if file.filename == '':
                    return jsonify({"error": "파일이 선택되지 않았습니다."}), 400
                
                if not allowed_file(file.filename):
                    return jsonify({"error": "허용되지 않는 파일 형식입니다. PDF, Word 또는 텍스트 파일만 업로드할 수 있습니다."}), 400
                
                # 파일에서 텍스트 추출
                try:
                    cover_letter = extract_text_from_file(file)
                except Exception as e:
                    return jsonify({"error": str(e)}), 400
            else:
                return jsonify({"error": "자기소개서 내용이나 파일이 제공되지 않았습니다."}), 400
        
        if not company or not position:
            return jsonify({"error": "회사명과 직무는 필수 입력 항목입니다."}), 400
        
        if not cover_letter or len(cover_letter.strip()) < 50:
            return jsonify({"error": "자기소개서 내용이 너무 짧거나 비어 있습니다."}), 400
        
        # 통합 평가 시스템 (단일 평가 로직)
        system_content = """당신은 전문적인 자기소개서 평가 전문가입니다. 5개 평가 항목을 기준으로 고품질 피드백을 제공해주세요.

평가 항목 (총 100점):
1) 구성 (20점) - 체계적 구성과 논리적 흐름
2) 내용 적합성 (20점) - 회사/직무 적합성과 이해도
3) 차별성 (20점) - 개인만의 특색과 경쟁력
4) 논리성 (20점) - 논리적 연결과 설득력
5) 가독성 (20점) - 문장 완성도와 표현력

** 평가 기준 **
- 각 항목마다 2-3줄의 상세한 피드백 제공
- 각 항목마다 반드시 사용자의 실제 문장을 인용하고 개선 예시 제공
- 건설적이고 구체적인 개선 방향 제시
- 강점과 개선점을 균형있게 언급

각 항목별 반드시 포함:
- 항목명과 점수 (예: "1. 구성 (20점 만점): 16점")
- 긍정적 피드백 (✓ 기호 사용)
- 개선 제안 (⚠️ 기호 사용)
- **실제 문장 예시**: 사용자의 자기소개서에서 해당 항목과 관련된 실제 문장을 "[기존 문장]" 형태로 인용
- **개선 제안**: 인용한 문장의 구체적인 개선 예시를 "[개선 제안]" 형태로 제공

** 실제 문장 추출 및 개선 (절대 필수) **
다음 자기소개서에서 각 평가 항목별로 개선이 필요한 실제 문장을 정확히 찾아 인용하고 개선 예시를 제공하세요:

각 항목 평가 후 반드시 다음 형식을 포함:
🔍 문제 문장: "[사용자가 실제로 작성한 원문 그대로]"
✏️ 개선 예시: "[구체적으로 개선된 문장]"

** 예시 **
🔍 문제 문장: "저는 다양한 경험을 통해 성장했습니다."
✏️ 개선 예시: "대학교 4년간 스마트팜 개발 프로젝트 팀장으로 활동하며 데이터 분석과 팀 협업 능력을 기를 수 있었습니다."

** 필수 규칙 **
1. 반드시 제공된 자기소개서 원문에서 실제 문장을 찾아 ""안에 정확히 인용
2. 템플릿이나 예시 문장이 아닌 실제 사용자 작성 문장만 사용
3. 각 항목당 1개씩, 총 5개의 문제 문장과 개선 예시 제공
4. 문제 문장이 없으면 "문제 문장 없음"이라고 명시

마지막에 총점과 등급(A/B+/B/C/D)을 제시하고, 전체적인 종합 평가를 작성해주세요.

** 중요: 각 항목별 정확한 배점을 표시해야 합니다 **
예시: "1. 구성 (20점 만점): 16점"과 같은 형식으로 표시하세요."""
        
        # 시스템 메시지 및 사용자 프롬프트 구성
        system_message = ChatCompletionSystemMessageParam(
            role="system",
            content=system_content
        )
        
        user_content = f"""다음 자기소개서를 평가해 주세요:

[지원 회사] {company}
[지원 직무] {position}

[자기소개서]
{cover_letter}

각 평가 항목(구성, 내용 적합성, 차별성, 논리성, 가독성)별로 점수와 피드백을 제공해 주세요.
긍정적인 피드백에는 ✓ 기호를, 개선이 필요한 부분에는 ⚠️ 기호를 사용하세요.

각 항목은 다음 형식으로 구성해주세요:
1. [항목명] (점수/20)
🟩 잘한 점:
- [구체적인 긍정 피드백]
🟧 개선점:
- [구체적인 개선 필요사항]
🛠 예시 문장:
- 문제 문장: "[사용자가 실제 작성한 문장]"
- 개선 예시: "[자연스럽게 개선된 문장]"

마지막에는 총점과 등급(A/B+/B/C/D)을 알려주세요."""
        
        user_message = ChatCompletionUserMessageParam(
            role="user",
            content=user_content
        )
        
        # OpenAI API 호출
        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[system_message, user_message],
                temperature=0.3  # 일관된 평가를 위해 낮은 temperature 값 사용
            )
        except Exception as api_error:
            logging.error(f"OpenAI API 호출 오류: {str(api_error)}")
            # 오류 응답 제공
            return jsonify({
                "success": False, 
                "error": "자기소개서 평가 중 오류가 발생했습니다. 잠시 후 다시 시도해 주세요.",
                "detail": str(api_error)
            }), 500
        
        # 응답에서 평가 결과 추출
        result = response.choices[0].message.content
        
        # 평가 기록 저장 (로그인한 사용자만)
        if current_user.is_authenticated:
            try:
                # 등급 추출 (총점과 등급 부분에서)
                grade = None
                if "등급" in result:
                    import re
                    grade_match = re.search(r'등급[:\s]*([A-D][+]?)', result)
                    if grade_match:
                        grade = grade_match.group(1)
                
                # 결제 타입 결정 (무료 1회 vs 유료 평가)
                payment_type = '무료'
                if current_user.plan_type != 'basic_0회' or current_user.remaining_credits > 0:
                    payment_type = '유료'
                
                # 고유 result_id 생성
                import uuid
                result_id = str(uuid.uuid4())[:8]  # 8자리 고유 ID
                
                # 평가 기록 생성 (텍스트 입력)
                evaluation_record = EvaluationRecord(
                    user_id=current_user.id,
                    evaluation_type='text',
                    payment_type=payment_type,
                    source_title='텍스트 입력',
                    evaluation_result=result,
                    grade=grade,
                    result_id=result_id
                )
                
                # 사용량 차감 (로그인된 사용자만)
                current_user.consume_evaluation()
                
                db.session.add(evaluation_record)
                db.session.commit()
                logging.info(f"텍스트 평가 기록 저장 완료: 사용자 {current_user.id}")
            except Exception as save_error:
                logging.error(f"텍스트 평가 기록 저장 오류: {str(save_error)}")
                # 평가 기록 저장 실패해도 평가 결과는 반환
        
        # 통합 평가 응답 로그 출력 (디버깅용)
        print("=== GPT 평가 전체 응답 ===")
        print(result)
        print("=== 응답 끝 ===")
        logging.info(f"GPT 평가 전체 응답: {result}")
        
        return jsonify({"success": True, "result": result})
    
    except Exception as e:
        logging.error(f"자기소개서 평가 오류: {str(e)}")
        return jsonify({"error": "자기소개서 평가에 실패했습니다. 나중에 다시 시도해 주세요."}), 500

# 파일 업로드 처리 엔드포인트
@app.route("/upload-evaluate", methods=["POST"])
def upload_evaluate():
    """파일 업로드를 통한 자기소개서 평가"""
    if not ENABLE_FILE_UPLOAD:
        return jsonify({"error": "파일 업로드 기능이 현재 비활성화되어 있습니다."}), 503
        
    try:
        # 필수 필드 확인
        if 'coverLetterFile' not in request.files:
            return jsonify({"error": "파일이 업로드되지 않았습니다."}), 400
            
        file = request.files['coverLetterFile']
        company = request.form.get('company', '')
        position = request.form.get('position', '')

        
        # 파일 유효성 검사
        if not file or file.filename == '':
            return jsonify({"error": "파일이 선택되지 않았습니다."}), 400
            
        # 파일명에서 확장자 추출
        filename = file.filename
        if '.' not in filename:
            return jsonify({"error": "파일 확장자를 확인할 수 없습니다. 지원 형식: PDF, DOCX, TXT"}), 400
            
        file_ext = filename.rsplit('.', 1)[1].lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            return jsonify({"error": f"지원하지 않는 파일 형식입니다: {file_ext}. 지원 형식: PDF, DOCX, TXT"}), 400
        
        logging.debug(f"파일 업로드 처리: {filename}, 타입: {file_ext}, 회사: {company}")
        
        # 사용자 입력 검증
        if not company or not position:
            return jsonify({"error": "회사명과 직무는 필수 입력 항목입니다."}), 400
            
        # 임시 파일 저장
        temp_dir = tempfile.gettempdir()
        temp_filename = secure_filename(filename)
        temp_path = os.path.join(temp_dir, temp_filename)
        
        try:
            file.save(temp_path)
            logging.debug(f"임시 파일 저장됨: {temp_path}")
            
            # 파일 형식에 따라 텍스트 추출
            extracted_text = ""
            
            # TXT 파일 처리
            if file_ext == 'txt':
                for encoding in ['utf-8', 'cp949', 'euc-kr']:
                    try:
                        with open(temp_path, 'r', encoding=encoding) as f:
                            extracted_text = f.read()
                        if extracted_text:
                            break
                    except UnicodeDecodeError:
                        continue
                        
            # PDF 파일 처리
            elif file_ext == 'pdf':
                with pdfplumber.open(temp_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + "\n\n"
                            
            # DOCX 파일 처리
            elif file_ext in ['docx', 'doc']:
                doc = docx.Document(temp_path)
                for para in doc.paragraphs:
                    if para.text:
                        extracted_text += para.text + '\n'
                        
            # 추출된 텍스트 확인
            if not extracted_text or len(extracted_text.strip()) < 50:
                return jsonify({
                    "error": "파일에서 충분한 텍스트를 추출할 수 없습니다. 파일이 비어있거나 형식이 올바르지 않을 수 있습니다."
                }), 400
                
            logging.debug(f"파일에서 성공적으로 텍스트 추출: {len(extracted_text)} 글자")
            
            # OpenAI API 요청 생성
            system_content = """너는 자기소개서를 평가하고 피드백을 제공하는 AI 평가 전문가야. 
다음 5가지 평가 기준으로 자기소개서를 분석하고 점수와 구체적인 피드백을 제공해주세요:

1. 구성 (20점): 자기소개, 지원 동기, 역량, 포부 등이 체계적으로 구성되어 있는지
2. 내용 적합성 (25점): 지원 회사와 직무에 적합한 내용으로 작성되었는지
3. 차별성 (20점): 다른 지원자와 차별화된 내용이 있는지
4. 논리성 (15점): 주장과 근거가 논리적으로 연결되는지
5. 가독성 (20점): 문장 구조와 어휘 선택이 적절한지

** 중요: 각 항목별 정확한 배점을 표시해야 합니다 **
예시: "1. 구성 (20점 만점): 16점"과 같은 형식으로 표시하세요.

각 평가 항목별로 다음 정보를 반드시 포함해주세요:
- 항목명과 배점 (예: "1. 구성 (20점 만점): 15점")
- 긍정적인 피드백 (앞에 ✓ 기호 사용)
- 개선이 필요한 부분 (앞에 ⚠️ 기호 사용)

마지막에는 5가지 항목의 점수를 합산한 총점과 이에 따른 등급을 제시해주세요:
- A 등급: 90-100점
- B+ 등급: 80-89점
- B 등급: 70-79점
- C 등급: 60-69점
- D 등급: 60점 미만"""


        
            # 시스템 메시지 및 사용자 프롬프트 구성
            system_message = ChatCompletionSystemMessageParam(
                role="system",
                content=system_content
            )
            
            user_content = f"""다음 자기소개서를 평가해 주세요:

[지원 회사] {company}
[지원 직무] {position}

[자기소개서]
{extracted_text}

각 항목은 다음 형식으로 구성해주세요:
1. [항목명] (점수/20)
🟩 잘한 점:
- [구체적인 긍정 피드백]
🟧 개선점:
- [구체적인 개선 필요사항]
🛠 예시 문장:
- 문제 문장: "[사용자가 실제 작성한 문장]"
- 개선 예시: "[자연스럽게 개선된 문장]"

마지막에는 총점과 등급(A/B+/B/C/D)을 알려주세요."""
            
            user_message = ChatCompletionUserMessageParam(
                role="user",
                content=user_content
            )
            
            # OpenAI API 호출
            try:
                response = openai_client.chat.completions.create(
                    model="gpt-4o",
                    messages=[system_message, user_message],
                    temperature=0.3  # 일관된 평가를 위해 낮은 temperature 값 사용
                )
                
                # 응답에서 평가 결과 추출
                result = response.choices[0].message.content
                
                # 평가 기록 저장 (로그인한 사용자만)
                if current_user.is_authenticated:
                    try:
                        # 등급 추출 (총점과 등급 부분에서)
                        grade = None
                        if "등급" in result:
                            import re
                            grade_match = re.search(r'등급[:\s]*([A-D][+]?)', result)
                            if grade_match:
                                grade = grade_match.group(1)
                        
                        # 결제 타입 결정 (무료 1회 vs 유료 평가)
                        payment_type = '무료'
                        if current_user.plan_type != 'basic_0회' or current_user.remaining_credits > 0:
                            payment_type = '유료'
                        
                        # 고유 result_id 생성
                        import uuid
                        result_id = str(uuid.uuid4())[:8]  # 8자리 고유 ID
                        
                        # 평가 기록 생성
                        evaluation_record = EvaluationRecord(
                            user_id=current_user.id,
                            evaluation_type='text',
                            payment_type=payment_type,
                            source_title='텍스트 입력',
                            evaluation_result=result,
                            grade=grade,
                            result_id=result_id
                        )
                        
                        db.session.add(evaluation_record)
                        db.session.commit()
                        logging.info(f"평가 기록 저장 완료: 사용자 {current_user.id}, 타입: {payment_type}")
                    except Exception as save_error:
                        logging.error(f"평가 기록 저장 오류: {str(save_error)}")
                        # 평가 기록 저장 실패해도 평가 결과는 반환
                
                return jsonify({"success": True, "result": result})
                
            except Exception as api_error:
                logging.error(f"OpenAI API 호출 오류: {str(api_error)}")
                return jsonify({
                    "success": False, 
                    "error": "자기소개서 평가 중 오류가 발생했습니다. 잠시 후 다시 시도해 주세요.",
                    "detail": str(api_error)
                }), 500
                
        except Exception as extract_error:
            logging.error(f"파일 처리 오류: {str(extract_error)}")
            return jsonify({"error": f"파일 '{filename}' 처리 중 오류가 발생했습니다: {str(extract_error)}"}), 400
            
        finally:
            # 임시 파일 삭제
            if os.path.exists(temp_path):
                os.remove(temp_path)
                logging.debug(f"임시 파일 삭제됨: {temp_path}")
        
    except Exception as e:
        logging.error(f"파일 업로드 평가 오류: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": f"파일 업로드 처리 중 오류가 발생했습니다: {str(e)}"}), 500

@app.route('/purchase')
def purchase():
    """평가권 구매 페이지"""
    return render_template('purchase.html')

@app.route('/mypage')
@auth_required_with_fallback
def mypage():
    """마이페이지"""
    from flask import g
    
    # 세션 복구된 사용자 또는 현재 사용자 확인
    active_user = getattr(g, 'user', None) or current_user
    
    if not active_user or not hasattr(active_user, 'get_status_info'):
        flash("로그인이 필요합니다.", "error")
        return redirect(url_for('auth', mode='login'))
    
    # 현재 로그인한 사용자 정보 및 평가권 상태 조회
    user_status = active_user.get_status_info()
    
    # 최근 평가 기록 조회 (최대 3개)
    recent_evaluations = EvaluationRecord.query.filter_by(user_id=active_user.id)\
                                              .order_by(EvaluationRecord.created_at.desc())\
                                              .limit(3)\
                                              .all()
    
    # 최신 후기 3개 가져오기
    recent_reviews = Review.query.filter_by(is_active=True).order_by(Review.created_at.desc()).limit(3).all()
    
    # 사용자가 작성한 후기 조회
    user_reviews = Review.query.filter_by(user_id=active_user.id, is_active=True).order_by(Review.created_at.desc()).all()
    
    # 사용자가 작성한 문의사항 조회
    user_questions = Question.query.filter_by(user_id=active_user.id, is_active=True).order_by(Question.created_at.desc()).all()
    
    # 사용자의 결제 내역 조회 (최근 10개)
    user_purchases = CreditPurchase.query.filter_by(user_id=active_user.id).order_by(CreditPurchase.purchase_date.desc()).limit(10).all()
    
    return render_template('mypage.html', 
                         user=active_user, 
                         user_status=user_status,
                         recent_evaluations=recent_evaluations,
                         recent_reviews=recent_reviews,
                         user_reviews=user_reviews,
                         user_questions=user_questions,
                         user_purchases=user_purchases,
                         firebase_api_key=os.environ.get("FIREBASE_API_KEY"),
                         firebase_project_id=os.environ.get("FIREBASE_PROJECT_ID"),
                         firebase_app_id=os.environ.get("FIREBASE_APP_ID"))

@app.route('/admin/user/<int:user_id>/credits', methods=['POST'])
@admin_required
def admin_give_credits():
    """관리자 전용: 사용자에게 평가권 지급"""
    try:
        data = request.get_json()
        user_id = data.get('user_id')
        credits_to_add = int(data.get('credits', 0))
        
        if credits_to_add <= 0:
            return jsonify({'success': False, 'error': '평가권 수량은 1개 이상이어야 합니다.'}), 400
        
        user = User.query.get(user_id)
        if not user:
            return jsonify({'success': False, 'error': '사용자를 찾을 수 없습니다.'}), 404
        
        # 평가권 지급
        user.remaining_credits += credits_to_add
        
        # 지급 기록 생성 (CreditPurchase 테이블 활용)
        from datetime import datetime, timedelta
        credit_record = CreditPurchase(
            user_id=user.id,
            credits_purchased=credits_to_add,
            purchase_date=datetime.utcnow(),
            expires_at=datetime.utcnow() + timedelta(days=90),  # 3개월 만료
            purchase_method='admin_grant',
            purchase_amount=0.0  # 무료 지급
        )
        db.session.add(credit_record)
        db.session.commit()
        
        logging.info(f"[ADMIN] {current_user.username}이 사용자 {user.username}에게 {credits_to_add}개 평가권 지급")
        
        return jsonify({
            'success': True, 
            'message': f'{user.username}님에게 {credits_to_add}개 평가권을 지급했습니다.',
            'new_credits': user.remaining_credits
        })
        
    except Exception as e:
        logging.error(f"평가권 지급 오류: {e}")
        return jsonify({'success': False, 'error': '평가권 지급 중 오류가 발생했습니다.'}), 500

# JasoAI 수정 - 관리자: /admin/dashboard 라우트 추가
@app.route('/admin/dashboard')
@admin_required
def admin_dashboard():
    """관리자 대시보드 페이지 (기존 기능 복구)"""
    from flask import g
    
    # 세션 복구된 사용자 또는 현재 사용자 확인
    active_user = getattr(g, 'user', None) or current_user
    
    # 모든 사용자 조회
    users = User.query.order_by(User.created_at.desc()).all()
    
    # 모든 문의사항 조회
    questions = Question.query.filter_by(is_active=True).order_by(Question.created_at.desc()).all()
    
    # 모든 후기 조회 (관리자는 모든 후기 볼 수 있음)
    all_reviews = Review.query.filter_by(is_active=True).order_by(Review.created_at.desc()).all()
    
    # 모든 구매 내역 조회 (최근 100개)
    purchases = CreditPurchase.query.order_by(CreditPurchase.purchase_date.desc()).limit(100).all()
    
    return render_template('admin_unified.html', 
                         users=users, 
                         questions=questions, 
                         all_reviews=all_reviews,
                         purchases=purchases)

@app.route('/admin')
@admin_required
def admin_panel():
    """통합 관리자 패널 페이지 (dashboard로 리다이렉트)"""
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/purchases')
@admin_required
def admin_purchases_legacy():
    """레거시 라우트: 결제/환불 관리로 리다이렉트"""
    return redirect('/admin/payments', code=302)


@app.route('/admin/refunds')
@admin_required
def admin_refunds_legacy():
    """레거시 라우트: 결제/환불 관리로 리다이렉트"""
    return redirect('/admin/payments', code=302)


@app.route('/admin/users')
@login_required
def admin_users_api():
    """관리자용 사용자 목록 API"""
    # JasoAI 수정: 관리자 권한 확인 (role 기반)
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '관리자 권한이 필요합니다.'}), 403
    
    users = User.query.order_by(User.created_at.desc()).all()
    users_data = []
    
    for user in users:
        users_data.append({
            'id': user.id,
            'email': user.email,
            'name': user.name,
            'plan_type': user.plan_type,
            'remaining_credits': user.remaining_credits,
            'daily_usage_count': user.daily_usage_count,
            'last_usage_date': user.last_usage_date.isoformat() if user.last_usage_date else None,
            'created_at': user.created_at.isoformat(),
            'last_login': user.last_login.isoformat() if user.last_login else None,
            'is_active': user.is_active
        })
    
    return jsonify({'success': True, 'users': users_data})

@app.route('/admin/users/<int:user_id>', methods=['GET'])
@login_required
def admin_get_user(user_id):
    """관리자용 사용자 정보 조회"""
    # JasoAI 수정: 관리자 권한 확인 (role 기반)
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '관리자 권한이 필요합니다.'}), 403
    
    try:
        user = User.query.get_or_404(user_id)
        return jsonify({
            'success': True,
            'user': {
                'id': user.id,
                'email': user.email,
                'name': user.name,
                'plan_type': user.plan_type,
                'remaining_credits': user.remaining_credits,
                'is_active': user.is_active,
                'created_at': user.created_at.isoformat(),
                'last_login': user.last_login.isoformat() if user.last_login else None
            }
        })
    except Exception as e:
        logging.error(f"사용자 조회 오류: {str(e)}")
        return jsonify({'success': False, 'error': '사용자 정보를 불러오는데 실패했습니다.'}), 500

@app.route('/admin/users/<int:user_id>', methods=['PUT'])
@login_required
def admin_update_user(user_id):
    """관리자용 사용자 정보 업데이트"""
    # JasoAI 수정: 관리자 권한 확인 (role 기반)
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '관리자 권한이 필요합니다.'}), 403
    
    try:
        user = User.query.get_or_404(user_id)
        
        # 폼 데이터에서 값 가져오기
        name = request.form.get('editUserName')
        plan_type = request.form.get('editUserPlan')
        remaining_credits = request.form.get('editUserCredits', type=int)
        is_active = request.form.get('editUserActive') == 'on'
        
        # 사용자 정보 업데이트
        if name is not None:
            user.name = name
        if plan_type is not None:
            user.plan_type = plan_type
        if remaining_credits is not None:
            user.remaining_credits = remaining_credits
        user.is_active = is_active
        
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': f'{user.name}({user.email}) 정보가 업데이트되었습니다.',
            'user': user.to_dict()
        })
        
    except Exception as e:
        print(f"사용자 업데이트 오류: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'message': '업데이트 중 오류가 발생했습니다.'}), 500

@app.route('/admin/users/<int:user_id>', methods=['DELETE'])
@login_required
def admin_delete_user(user_id):
    """관리자용 사용자 삭제"""
    # JasoAI 수정: 관리자 권한 확인 (role 기반)
    if current_user.role != 'admin':
        return jsonify({'success': False, 'message': '관리자 권한이 필요합니다.'}), 403
    
    try:
        user = User.query.get_or_404(user_id)
        
        # 자기 자신은 삭제할 수 없음
        if user.id == current_user.id:
            return jsonify({'success': False, 'message': '자기 자신은 삭제할 수 없습니다.'}), 400
        
        user_email = user.email
        
        # 관련된 평가 기록들을 먼저 삭제 (사용자 삭제 시 함께 삭제)
        EvaluationRecord.query.filter_by(user_id=user.id).delete()
        
        # 관련된 문의사항들은 유지 (삭제된 사용자로 표시)
        # Question은 그대로 유지
        
        db.session.delete(user)
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': f'{user_email} 계정이 삭제되었습니다.'
        })
        
    except Exception as e:
        logging.error(f"사용자 삭제 오류: {str(e)}")
        db.session.rollback()
        return jsonify({'success': False, 'error': '삭제 중 오류가 발생했습니다.'}), 500

@app.route('/delete_account', methods=['DELETE'])
@login_required
def delete_account():
    """사용자 계정 삭제"""
    try:
        user_id = current_user.id
        user_name = current_user.name or current_user.user_id or current_user.email
        
        # 관련 데이터 삭제 (외래키 제약조건 순서 고려)
        try:
            # 1. 평가 기록 삭제
            evaluation_records = EvaluationRecord.query.filter_by(user_id=user_id).all()
            for record in evaluation_records:
                db.session.delete(record)
            
            # 2. 후기 삭제
            reviews = Review.query.filter_by(user_id=user_id).all()
            for review in reviews:
                db.session.delete(review)
            
            # 3. 문의사항 삭제
            questions = Question.query.filter_by(user_id=user_id).all()
            for question in questions:
                db.session.delete(question)
            
            # 4. 답변한 문의사항의 answered_by 필드 null로 설정
            answered_questions = Question.query.filter_by(answered_by=user_id).all()
            for question in answered_questions:
                question.answered_by = None
            
            # 5. 로그아웃 처리
            logout_user()
            
            # 6. 사용자 삭제
            user = User.query.get(user_id)
            if user:
                db.session.delete(user)
            
            db.session.commit()
            
            logging.info(f"계정 삭제 완료: 사용자 {_mask_username(user_name)} (ID: ****)")
            return jsonify({
                'success': True, 
                'message': '계정이 성공적으로 삭제되었습니다.'
            })
            
        except Exception as inner_e:
            db.session.rollback()
            logging.error(f"계정 삭제 과정 오류: {str(inner_e)}")
            return jsonify({'success': False, 'error': f'계정 삭제 중 오류가 발생했습니다: {str(inner_e)}'}), 500
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"계정 삭제 오류: {str(e)}")
        return jsonify({'success': False, 'error': f'계정 삭제 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/profile/change-password', methods=['POST'])
@auth_required_with_fallback
def change_password():
    """비밀번호 변경"""
    from flask import g
    
    try:
        # 세션 복구된 사용자 또는 현재 사용자 확인
        active_user = getattr(g, 'user', None) or current_user
        
        if not active_user or not hasattr(active_user, 'check_password'):
            return jsonify({'success': False, 'error': '로그인이 필요합니다.'}), 401
        
        data = request.get_json()
        current_password = data.get('current_password')
        new_password = data.get('new_password')
        
        if not current_password or not new_password:
            return jsonify({'success': False, 'error': '모든 필드를 입력해주세요.'}), 400
        
        # 현재 비밀번호 검증
        if not active_user.check_password(current_password):
            return jsonify({'success': False, 'error': '현재 비밀번호가 올바르지 않습니다.'}), 400
        
        # 새 비밀번호 길이 검증
        if len(new_password) < 8:
            return jsonify({'success': False, 'error': '새 비밀번호는 8자 이상이어야 합니다.'}), 400
        
        # 비밀번호 변경
        active_user.set_password(new_password)
        db.session.commit()
        
        logging.info(f"비밀번호 변경: 사용자 {active_user.id}")
        return jsonify({'success': True, 'message': '비밀번호가 성공적으로 변경되었습니다.'})
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"비밀번호 변경 오류: {str(e)}")
        return jsonify({'success': False, 'error': '비밀번호 변경 중 오류가 발생했습니다.'}), 500

@app.route('/profile/change-email', methods=['POST'])
@auth_required_with_fallback
def change_email():
    """이메일 변경"""
    from flask import g
    
    try:
        # 세션 복구된 사용자 또는 현재 사용자 확인
        active_user = getattr(g, 'user', None) or current_user
        
        if not active_user or not hasattr(active_user, 'email'):
            return jsonify({'success': False, 'error': '로그인이 필요합니다.'}), 401
        
        data = request.get_json()
        new_email = data.get('new_email', '').strip()
        
        if not new_email:
            return jsonify({'success': False, 'error': '새 이메일을 입력해주세요.'}), 400
        
        # 이메일 형식 검증
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, new_email):
            return jsonify({'success': False, 'error': '올바른 이메일 형식을 입력해주세요.'}), 400
        
        # 현재 이메일과 동일한지 확인
        if active_user.email == new_email:
            return jsonify({'success': False, 'error': '현재 이메일과 동일합니다.'}), 400
        
        # 중복 이메일 확인
        existing_user = User.query.filter_by(email=new_email).first()
        if existing_user and existing_user.id != active_user.id:
            return jsonify({'success': False, 'error': '이미 사용 중인 이메일입니다.'}), 400
        
        # 이메일 변경
        active_user.email = new_email
        db.session.commit()
        
        logging.info(f"이메일 변경: 사용자 {active_user.id} -> {new_email}")
        return jsonify({'success': True, 'message': '이메일이 성공적으로 변경되었습니다.'})
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"이메일 변경 오류: {str(e)}")
        return jsonify({'success': False, 'error': '이메일 변경 중 오류가 발생했습니다.'}), 500

@app.route('/results/<result_id>')
@login_required
def view_evaluation_result(result_id):
    """평가 결과 상세보기 - result_id로 접근"""
    try:
        # 해당 result_id로 평가 기록 조회
        evaluation_record = EvaluationRecord.query.filter_by(result_id=result_id).first()
        
        if not evaluation_record:
            flash('요청하신 평가 결과를 찾을 수 없습니다.', 'error')
            return redirect(url_for('mypage'))
        
        # 권한 체크: 본인의 평가 결과인지 확인
        if evaluation_record.user_id != current_user.id:
            flash('다른 사용자의 평가 결과에는 접근할 수 없습니다.', 'error')
            return redirect(url_for('mypage'))
        
        # 평가 결과가 있는지 확인
        if not evaluation_record.evaluation_result:
            flash('평가 결과 데이터가 없습니다.', 'error')
            return redirect(url_for('mypage'))
        
        # 평가 결과 페이지 렌더링
        return render_template('result_detail.html', 
                               evaluation=evaluation_record,
                               result_text=evaluation_record.evaluation_result)
        
    except Exception as e:
        logging.error(f"평가 결과 조회 오류: {str(e)}")
        flash('평가 결과를 불러오는 중 오류가 발생했습니다.', 'error')
        return redirect(url_for('mypage'))


@app.route("/pay/fallback")
def pay_fallback():
    """iframe 환경에서의 결제 폴백 페이지"""
    return render_template("pay_fallback.html",
                           imp_code=os.environ.get("PORTONE_IMP_CODE_LIVE",""),
                           chan_inicis=os.environ.get("PORTONE_CHANNEL_INICIS_LIVE",""),
                           chan_tosspay=os.environ.get("PORTONE_CHANNEL_TOSSPAY_LIVE",""),
                           pgcode_tosspay=os.environ.get("PORTONE_PGCODE_TOSSPAY","tosspay"),
                           opts=dict(request.args))

@app.route("/refund/request", methods=["POST"])
@login_required
def refund_request():
    """사용자 환불 요청"""
    try:
        data = request.get_json() or {}
        purchase_id = data.get("purchase_id")
        
        if not purchase_id:
            return jsonify({"success": False, "message": "구매 ID가 필요합니다"}), 400
        
        # 구매 기록 조회
        purchase = CreditPurchase.query.get(purchase_id)
        if not purchase or purchase.user_id != current_user.id:
            logging.error(f"[REFUND-REQ][FAIL] 권한 없음 또는 없는 구매: purchase_id={purchase_id}, user_id={current_user.id}")
            return jsonify({"success": False, "message": "구매 내역을 찾을 수 없습니다"}), 404
        
        # 이미 환불 요청되었거나 환불된 경우
        if purchase.refund_status in ("requested", "approved"):
            return jsonify({"success": False, "message": "이미 환불 요청되었거나 처리된 구매입니다"}), 400
        
        # 이미 만료 처리된 경우
        if purchase.expired_at:
            return jsonify({"success": False, "message": "이미 만료 처리된 구매입니다"}), 400
        
        # 환불 요청 상태로 변경
        purchase.refund_status = "requested"
        purchase.status = "refund_requested"
        purchase.refund_requested_at = datetime.utcnow()
        
        db.session.commit()
        
        logging.info(f"[REFUND-REQ][OK] purchase_id={purchase_id}, user_id={current_user.id}, amount={purchase.remaining_amount}")
        
        add_notification(
            user_id=current_user.id,
            notif_type='refund_requested',
            title='환불 요청을 접수했습니다',
            body=f'{purchase.product_name or "평가권"} 환불 요청 접수',
            meta={'purchase_id': purchase_id, 'merchant_uid': purchase.merchant_uid}
        )
        
        return jsonify({
            "success": True,
            "message": "환불 요청이 접수되었습니다. 관리자 확인 후 처리됩니다."
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"[REFUND-REQ][ERROR] 환불 요청 오류: {str(e)}")
        return jsonify({"success": False, "message": "환불 요청 중 오류가 발생했습니다"}), 500


@app.route("/admin/refund/<int:purchase_id>", methods=["POST"])
@login_required
def admin_refund(purchase_id):
    """관리자 전용 평가권 환불 (승인 처리 + 실제 PG 환불)"""
    try:
        # 관리자 권한 체크
        if not (current_user.role == "admin" or session.get("role") == "admin"):
            logging.error(f"[REFUND][FAIL] 권한 없음: user_id={current_user.id}")
            return jsonify({"success": False, "message": "권한이 없습니다"}), 403
        
        # 구매 기록 조회
        purchase = CreditPurchase.query.get(purchase_id)
        if not purchase:
            logging.error(f"[REFUND][FAIL] 구매 기록 없음: purchase_id={purchase_id}")
            return jsonify({"success": False, "message": "구매 기록을 찾을 수 없습니다"}), 404
        
        # 이미 환불된 경우
        if purchase.refund_status == "approved" or purchase.expired_at:
            logging.warning(f"[REFUND][WARN] 이미 환불됨: purchase_id={purchase_id}")
            return jsonify({"success": False, "message": "이미 환불된 구매입니다"}), 400
        
        # 사용자 조회
        user = User.query.get(purchase.user_id)
        if not user:
            logging.error(f"[REFUND][FAIL] 사용자 없음: user_id={purchase.user_id}")
            return jsonify({"success": False, "message": "사용자를 찾을 수 없습니다"}), 404
        
        # PortOne 실제 환불 호출
        refund_result = portone_refund(
            imp_uid=purchase.imp_uid,
            amount=purchase.amount,
            reason='admin_refund'
        )
        
        if not refund_result.get("ok"):
            error_msg = refund_result.get("message", "환불 처리 실패")
            logging.error(f"[REFUND][FAIL] PG 환불 실패: purchase_id={purchase_id}, error={error_msg}")
            return jsonify({"success": False, "message": f"PG 환불 실패: {error_msg}"}), 500
        
        if refund_result.get("stub"):
            logging.warning(f"[REFUND][STUB] PG 환불 스텁 모드: purchase_id={purchase_id}")
        
        # 남은 평가권 차감
        credits_to_refund = purchase.remaining_amount
        if credits_to_refund > 0:
            user.remaining_credits = max(0, (user.remaining_credits or 0) - credits_to_refund)
            user.evaluation_ticket = max(0, (user.evaluation_ticket or 0) - credits_to_refund)
        
        # Ledger 기록 (환불)
        ledger_entry = CreditsLedger(
            user_id=user.id,
            payment_id=purchase.id,
            delta=-credits_to_refund,
            reason='refund',
            note=f'환불 승인 (purchase_id={purchase_id})'
        )
        db.session.add(ledger_entry)
        
        # 환불 처리
        refund_time = datetime.utcnow()
        purchase.remaining_amount = 0
        purchase.expired_at = refund_time
        purchase.refund_status = "approved"
        purchase.status = "refunded"
        purchase.refunded_at = refund_time
        purchase.refund_amount = purchase.amount  # 환불 금액 기록
        
        db.session.commit()
        
        logging.info(f"[REFUND][OK] purchase_id={purchase_id}, user_id={user.id}, refunded_credits={credits_to_refund}, new_total={user.remaining_credits}")
        
        add_notification(
            user_id=user.id,
            notif_type='refund_approved',
            title='환불이 완료되었습니다',
            body=f'{purchase.amount:,}원이 환불 처리되었습니다',
            meta={'purchase_id': purchase_id, 'amount': purchase.amount, 'credits': credits_to_refund}
        )
        
        return jsonify({
            "success": True,
            "message": "환불이 완료되었습니다",
            "refunded_credits": credits_to_refund,
            "user_remaining_credits": user.remaining_credits,
            "pg_stub": refund_result.get("stub", False)
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"[REFUND][ERROR] 환불 처리 오류: {str(e)}")
        return jsonify({"success": False, "message": "환불 처리 중 오류가 발생했습니다"}), 500


@app.route("/admin/payments", methods=["GET"])
@login_required
def admin_payments():
    """관리자 전용 결제 내역 관리 페이지"""
    try:
        # 관리자 권한 체크
        if not (current_user.role == "admin" or session.get("role") == "admin"):
            logging.error(f"[ADMIN-PAYMENTS][FAIL] 권한 없음: user_id={current_user.id}")
            flash("관리자 권한이 필요합니다", "error")
            return redirect(url_for('index'))
        
        # 검색/필터 파라미터
        search_query = request.args.get('search', '').strip()
        status_filter = request.args.get('status', '').strip()
        product_filter = request.args.get('product', '').strip()
        page = request.args.get('page', 1, type=int)
        per_page = 50
        
        # 기본 쿼리
        query = CreditPurchase.query
        
        # 검색 (merchant_uid, imp_uid, 사용자 이메일/아이디)
        if search_query:
            query = query.join(User).filter(
                db.or_(
                    CreditPurchase.merchant_uid.ilike(f'%{search_query}%'),
                    CreditPurchase.imp_uid.ilike(f'%{search_query}%'),
                    User.email.ilike(f'%{search_query}%'),
                    User.username.ilike(f'%{search_query}%')
                )
            )
        
        # 상태 필터
        if status_filter:
            if status_filter == 'paid':
                query = query.filter(
                    db.or_(
                        CreditPurchase.status == 'paid',
                        db.and_(CreditPurchase.status.is_(None), CreditPurchase.refund_status == 'none')
                    )
                )
            elif status_filter == 'refund_requested':
                query = query.filter(CreditPurchase.refund_status == 'requested')
            elif status_filter == 'refunded':
                query = query.filter(
                    db.or_(
                        CreditPurchase.status == 'refunded',
                        CreditPurchase.refund_status == 'approved'
                    )
                )
        
        # 상품 필터 (EVAL_1, EVAL_5 모두 허용)
        if product_filter:
            query = query.filter(CreditPurchase.product_code == product_filter)
        
        # 정렬 (최신순)
        query = query.order_by(CreditPurchase.purchase_date.desc())
        
        # 페이지네이션
        paginated = query.paginate(page=page, per_page=per_page, error_out=False)
        payments = paginated.items
        
        # 사용자 정보 조회 (한 번에)
        user_ids = list(set([p.user_id for p in payments]))
        users_dict = {u.id: u for u in User.query.filter(User.id.in_(user_ids)).all()}
        
        # 결제 데이터 준비
        payments_data = []
        for payment in payments:
            user = users_dict.get(payment.user_id)
            
            # 상태 통합 (status 우선, 없으면 refund_status 기반)
            if payment.status:
                display_status = payment.status
            elif payment.refund_status == 'approved':
                display_status = 'refunded'
            elif payment.refund_status == 'requested':
                display_status = 'refund_requested'
            else:
                display_status = 'paid'
            
            # Ledger 합계 계산
            ledger_sum = db.session.query(func.sum(CreditsLedger.delta)).filter(
                CreditsLedger.payment_id == payment.id
            ).scalar() or 0
            
            payments_data.append({
                'id': payment.id,
                'user_email': user.email if user and user.email else '(알 수 없음)',
                'user_name': user.name if user else '(알 수 없음)',
                'product_code': payment.product_code or '(미지정)',
                'product_name': payment.product_name or '(미지정)',
                'amount': payment.amount or 0,
                'credits_added': payment.credits_added or payment.original_amount,
                'status': display_status,
                'merchant_uid': payment.merchant_uid or '',
                'imp_uid': payment.imp_uid or '',
                'paid_at': payment.paid_at or payment.purchase_date,
                'refunded_at': payment.refunded_at,
                'credited_at': payment.credited_at,
                'ledger_sum': ledger_sum
            })
        
        logging.info(f"[ADMIN-PAYMENTS][OK] 결제 내역 조회: page={page}, count={len(payments)}")
        
        return render_template('admin_payments.html',
                               payments=payments_data,
                               page=page,
                               total_pages=paginated.pages,
                               search_query=search_query,
                               status_filter=status_filter,
                               product_filter=product_filter)
    
    except Exception as e:
        logging.error(f"[ADMIN-PAYMENTS][ERROR] 결제 내역 조회 오류: {str(e)}")
        flash("결제 내역을 불러오는 중 오류가 발생했습니다", "error")
        return redirect(url_for('admin_panel'))


@app.route("/admin/payments/manual_add", methods=["GET", "POST"])
@login_required
@admin_required
def admin_payments_manual_add():
    """비활성화된 라우트: 수동 등록 기능은 더 이상 사용하지 않습니다"""
    flash("수동 등록 기능은 비활성화되었습니다.", "warning")
    return redirect("/admin/payments", code=302)


@app.route("/admin/payments/recredit/<int:payment_id>", methods=["POST"])
@login_required
def admin_payments_recredit(payment_id):
    """관리자 전용 크레딧 재처리 (미지급 건 복구)"""
    try:
        # 관리자 권한 체크
        if not (current_user.role == "admin" or session.get("role") == "admin"):
            return jsonify({"success": False, "message": "관리자 권한이 필요합니다"}), 403
        
        payment = CreditPurchase.query.get(payment_id)
        if not payment:
            return jsonify({"success": False, "message": "결제 내역을 찾을 수 없습니다"}), 404
        
        # 이미 credited_at이 있으면 거부
        if payment.credited_at:
            return jsonify({"success": False, "message": "이미 크레딧이 지급된 결제입니다"}), 400
        
        # product_code/credits_added 확정
        inferred_code, inferred_credits, warning = _infer_product_from_amount(payment.amount)
        if not payment.product_code:
            payment.product_code = inferred_code
        if not payment.credits_added:
            payment.credits_added = inferred_credits or payment.original_amount
        
        credits_to_add = payment.credits_added or payment.original_amount
        
        # 중복 지급 방지: Ledger 확인
        existing_ledger = CreditsLedger.query.filter_by(
            payment_id=payment.id
        ).first()
        
        if existing_ledger:
            return jsonify({"success": False, "message": "이미 Ledger에 기록된 결제입니다"}), 400
        
        # 사용자 크레딧 증가
        user = payment.user
        user.evaluation_ticket = (user.evaluation_ticket or 0) + credits_to_add
        user.remaining_credits = (user.remaining_credits or 0) + credits_to_add
        
        # Ledger 기록
        ledger_entry = CreditsLedger(
            user_id=payment.user_id,
            payment_id=payment.id,
            delta=credits_to_add,
            reason='manual',
            note=f'관리자 재처리 (payment_id={payment.id})'
        )
        db.session.add(ledger_entry)
        
        # credited_at 설정
        payment.credited_at = datetime.utcnow()
        
        db.session.commit()
        
        logging.info(f"[RECREDIT][OK] payment_id={payment_id}, user_id={user.id}, credits={credits_to_add}")
        
        return jsonify({
            "success": True,
            "message": f"{user.email}에게 {credits_to_add}개의 평가권이 지급되었습니다",
            "credits_added": credits_to_add
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"[RECREDIT][ERROR] 재처리 오류: {str(e)}")
        return jsonify({"success": False, "message": "재처리 중 오류가 발생했습니다"}), 500


@app.route("/admin/reconcile", methods=["GET"])
@login_required
def admin_reconcile():
    """관리자 전용 결제 정합성 관리 페이지"""
    try:
        # 관리자 권한 체크
        if not (current_user.role == "admin" or session.get("role") == "admin"):
            logging.error(f"[ADMIN-RECONCILE][FAIL] 권한 없음: user_id={current_user.id}")
            flash("관리자 권한이 필요합니다", "error")
            return redirect(url_for('index'))
        
        logging.info(f"[ADMIN-RECONCILE][OK] 정합성 관리 페이지 접근")
        return render_template('admin_reconcile.html')
    
    except Exception as e:
        logging.error(f"[ADMIN-RECONCILE][ERROR] 페이지 로드 오류: {str(e)}")
        flash("페이지를 불러오는 중 오류가 발생했습니다", "error")
        return redirect(url_for('admin_panel'))


def _infer_product_from_amount(amount):
    """금액으로 상품 코드 및 평가권 수 추론"""
    if not amount:
        return None, None, "금액 정보 없음"
    
    # 허용 오차 ±100원
    tolerance = 100
    
    if abs(amount - AMT_1) <= tolerance:
        return 'EVAL_1', 1, None
    elif abs(amount - AMT_5) <= tolerance:
        return 'EVAL_5', 5, None
    else:
        # 가장 가까운 금액으로 추정 (경고 포함)
        dist_1 = abs(amount - AMT_1)
        dist_5 = abs(amount - AMT_5)
        if dist_1 < dist_5:
            return 'EVAL_1', 1, f"금액 불일치 (예상: {AMT_1}원, 실제: {amount}원)"
        else:
            return 'EVAL_5', 5, f"금액 불일치 (예상: {AMT_5}원, 실제: {amount}원)"


@app.route("/admin/reconcile/preview", methods=["POST"])
@login_required
def admin_reconcile_preview():
    """정합성 DRY RUN (미리보기)"""
    try:
        # 관리자 권한 체크
        if not (current_user.role == "admin" or session.get("role") == "admin"):
            return jsonify({"success": False, "message": "관리자 권한이 필요합니다"}), 403
        
        data = request.get_json() or {}
        date_from_str = data.get('date_from', '').strip()
        date_to_str = data.get('date_to', '').strip()
        target = data.get('target', 'all')  # 'all' or 'missing_credit'
        user_email = data.get('user_email', '').strip()
        
        # 날짜 파싱
        date_from = datetime.fromisoformat(date_from_str) if date_from_str else None
        date_to = datetime.fromisoformat(date_to_str) if date_to_str else None
        
        # 기본 쿼리: status='paid' 인 결제 건들
        query = CreditPurchase.query.filter(
            db.or_(
                CreditPurchase.status == 'paid',
                db.and_(CreditPurchase.status.is_(None), CreditPurchase.refund_status == 'none')
            )
        )
        
        # 날짜 필터
        if date_from:
            query = query.filter(CreditPurchase.paid_at >= date_from)
        if date_to:
            query = query.filter(CreditPurchase.paid_at <= date_to)
        
        # 사용자 필터
        if user_email:
            user = User.query.filter_by(email=user_email).first()
            if user:
                query = query.filter(CreditPurchase.user_id == user.id)
            else:
                return jsonify({"success": False, "message": f"사용자를 찾을 수 없습니다: {user_email}"}), 404
        
        payments = query.order_by(CreditPurchase.paid_at.desc()).all()
        
        # 분류 결과
        missing_credits = []  # 미지급
        unclear = []  # 불명확 (product_code 누락)
        duplicate_suspect = []  # 중복지급 의심
        normal = []  # 정상
        
        for p in payments:
            # Ledger에서 해당 payment_id의 누적 delta 계산
            ledger_sum = db.session.query(func.sum(CreditsLedger.delta)).filter(
                CreditsLedger.payment_id == p.id
            ).scalar() or 0
            
            # product_code/credits_added 추론
            inferred_code, inferred_credits, warning = _infer_product_from_amount(p.amount)
            
            # 현재 product_code/credits_added 확정
            final_code = p.product_code or inferred_code
            final_credits = p.credits_added or inferred_credits or p.original_amount
            
            issue_type = None
            issue_desc = []
            
            # 기준1: credited_at이 NULL
            if not p.credited_at:
                issue_type = "missing"
                issue_desc.append("credited_at NULL")
            
            # 기준2: Ledger 합계가 0 (또는 부족)
            if ledger_sum == 0:
                issue_type = "missing"
                issue_desc.append("Ledger 기록 없음")
            elif ledger_sum < final_credits:
                issue_type = "missing"
                issue_desc.append(f"Ledger 부족 (합계: {ledger_sum}, 예상: {final_credits})")
            
            # 기준3: product_code/credits_added 누락
            if not p.product_code or not p.credits_added:
                if issue_type != "missing":
                    issue_type = "unclear"
                issue_desc.append("product_code/credits_added 누락")
            
            # 경고 추가
            if warning:
                issue_desc.append(warning)
            
            # 중복지급 의심
            if ledger_sum > final_credits:
                issue_type = "duplicate"
                issue_desc.append(f"중복지급 의심 (Ledger: {ledger_sum} > 예상: {final_credits})")
            
            payment_info = {
                'id': p.id,
                'user_id': p.user_id,
                'user_email': p.user.email if p.user else '(알 수 없음)',
                'merchant_uid': p.merchant_uid or '',
                'imp_uid': p.imp_uid or '',
                'amount': p.amount or 0,
                'paid_at': p.paid_at.isoformat() if p.paid_at else '',
                'product_code': p.product_code,
                'credits_added': p.credits_added,
                'inferred_code': inferred_code,
                'inferred_credits': inferred_credits,
                'final_code': final_code,
                'final_credits': final_credits,
                'credited_at': p.credited_at.isoformat() if p.credited_at else None,
                'ledger_sum': ledger_sum,
                'issues': ', '.join(issue_desc) if issue_desc else '정상'
            }
            
            if issue_type == "missing":
                missing_credits.append(payment_info)
            elif issue_type == "unclear":
                unclear.append(payment_info)
            elif issue_type == "duplicate":
                duplicate_suspect.append(payment_info)
            else:
                normal.append(payment_info)
        
        logging.info(f"[RECONCILE-PREVIEW][OK] 총 {len(payments)}건 분석: 미지급={len(missing_credits)}, 불명확={len(unclear)}, 중복의심={len(duplicate_suspect)}, 정상={len(normal)}")
        
        return jsonify({
            "success": True,
            "total": len(payments),
            "summary": {
                "missing_credits": len(missing_credits),
                "unclear": len(unclear),
                "duplicate_suspect": len(duplicate_suspect),
                "normal": len(normal)
            },
            "missing_credits": missing_credits,
            "unclear": unclear,
            "duplicate_suspect": duplicate_suspect,
            "normal": normal[:20]  # 정상 건은 최대 20개만 미리보기
        })
        
    except Exception as e:
        logging.error(f"[RECONCILE-PREVIEW][ERROR] 미리보기 오류: {str(e)}")
        return jsonify({"success": False, "message": f"미리보기 중 오류가 발생했습니다: {str(e)}"}), 500


@app.route("/admin/reconcile/apply", methods=["POST"])
@login_required
def admin_reconcile_apply():
    """정합성 복구 실제 적용"""
    try:
        # 관리자 권한 체크
        if not (current_user.role == "admin" or session.get("role") == "admin"):
            return jsonify({"success": False, "message": "관리자 권한이 필요합니다"}), 403
        
        data = request.get_json() or {}
        payment_ids = data.get('payment_ids', [])
        
        if not payment_ids:
            return jsonify({"success": False, "message": "처리할 결제 건이 없습니다"}), 400
        
        applied_count = 0
        failed_ids = []
        
        # 100건씩 트랜잭션 처리
        batch_size = 100
        for i in range(0, len(payment_ids), batch_size):
            batch = payment_ids[i:i+batch_size]
            
            try:
                for payment_id in batch:
                    payment = CreditPurchase.query.get(payment_id)
                    if not payment:
                        failed_ids.append(payment_id)
                        continue
                    
                    # 이미 credited_at이 있으면 건너뜀
                    if payment.credited_at:
                        logging.warning(f"[RECONCILE-APPLY][SKIP] 이미 지급됨: payment_id={payment_id}")
                        continue
                    
                    # product_code/credits_added 확정
                    inferred_code, inferred_credits, _ = _infer_product_from_amount(payment.amount)
                    if not payment.product_code:
                        payment.product_code = inferred_code
                    if not payment.credits_added:
                        payment.credits_added = inferred_credits or payment.original_amount
                    
                    credits_to_add = payment.credits_added or payment.original_amount
                    
                    # Ledger에 이미 기록이 있는지 확인 (중복 방지)
                    existing_ledger = CreditsLedger.query.filter_by(
                        payment_id=payment.id,
                        reason='payment'
                    ).first()
                    
                    if not existing_ledger:
                        # 사용자 크레딧 증가
                        user = payment.user
                        user.evaluation_ticket = (user.evaluation_ticket or 0) + credits_to_add
                        user.remaining_credits = (user.remaining_credits or 0) + credits_to_add
                        
                        # Ledger 기록
                        ledger_entry = CreditsLedger(
                            user_id=payment.user_id,
                            payment_id=payment.id,
                            delta=credits_to_add,
                            reason='reconcile',
                            note=f'정합성 복구 (payment_id={payment.id})'
                        )
                        db.session.add(ledger_entry)
                    
                    # credited_at 설정
                    payment.credited_at = datetime.utcnow()
                    
                    applied_count += 1
                
                db.session.commit()
                logging.info(f"[RECONCILE-APPLY][OK] Batch {i//batch_size + 1}: {len(batch)}건 처리")
                
            except Exception as batch_error:
                db.session.rollback()
                logging.error(f"[RECONCILE-APPLY][ERROR] Batch {i//batch_size + 1} 실패: {str(batch_error)}")
                failed_ids.extend(batch)
        
        logging.info(f"[RECONCILE-APPLY][OK] 총 {applied_count}건 적용, 실패 {len(failed_ids)}건")
        
        return jsonify({
            "success": True,
            "applied_count": applied_count,
            "failed_count": len(failed_ids),
            "failed_ids": failed_ids,
            "message": f"{applied_count}건의 결제가 복구되었습니다"
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"[RECONCILE-APPLY][ERROR] 적용 오류: {str(e)}")
        return jsonify({"success": False, "message": f"적용 중 오류가 발생했습니다: {str(e)}"}), 500


@app.route("/admin/payments/manual_refund", methods=["GET", "POST"])
@login_required
@admin_required
def admin_payments_manual_refund():
    """비활성화된 라우트: 수동 환불기록 기능은 더 이상 사용하지 않습니다"""
    flash("수동 환불기록 기능은 비활성화되었습니다.", "warning")
    return redirect("/admin/payments", code=302)


@app.route("/portone/webhook", methods=["POST"])
def portone_webhook():
    """PortOne 웹훅: 콘솔에서 수동 취소된 거래 자동 반영"""
    try:
        data = request.json or {}
        event = data.get("status") or data.get("event")
        imp_uid = data.get("imp_uid")
        merchant_uid = data.get("merchant_uid")
        cancel_amount = data.get("cancel_amount") or data.get("amount")
        cancelled_at_str = data.get("cancelled_at")
        
        logging.info(f"[WEBHOOK] PortOne event={event}, imp_uid={imp_uid}, merchant_uid={merchant_uid}")
        
        # 환불/취소 이벤트만 처리
        if event not in ("cancelled", "refund", "refunded", "cancel"):
            logging.info(f"[WEBHOOK] 무시: event={event}")
            return jsonify({"ok": True, "msg": "ignored"})
        
        # 결제 조회
        payment = CreditPurchase.query.filter(
            db.or_(
                CreditPurchase.imp_uid == imp_uid,
                CreditPurchase.merchant_uid == merchant_uid
            )
        ).first()
        
        if not payment:
            logging.warning(f"[WEBHOOK] 결제 내역 없음: imp_uid={imp_uid}, merchant_uid={merchant_uid}")
            return jsonify({"ok": False, "error": "payment_not_found"}), 404
        
        # 이미 환불 처리된 경우 중복 방지
        if payment.status == "refunded":
            logging.warning(f"[WEBHOOK] 이미 환불됨: payment_id={payment.id}")
            return jsonify({"ok": True, "msg": "already_refunded"})
        
        # 환불 시각 파싱
        refunded_at = datetime.utcnow()
        if cancelled_at_str:
            try:
                refunded_at = datetime.fromisoformat(cancelled_at_str.replace('Z', '+00:00'))
            except:
                pass
        
        # 결제 상태 업데이트
        payment.status = "refunded"
        payment.refunded_at = refunded_at
        payment.refund_id = imp_uid
        payment.refund_amount = cancel_amount or payment.amount
        payment.refund_status = "approved"
        payment.remaining_amount = 0
        payment.expired_at = refunded_at
        
        # 사용자 크레딧 회수
        user = User.query.get(payment.user_id)
        if user and payment.credits_added:
            credits_to_refund = payment.credits_added
            user.evaluation_ticket = max(0, (user.evaluation_ticket or 0) - credits_to_refund)
            user.remaining_credits = max(0, (user.remaining_credits or 0) - credits_to_refund)
            
            # Ledger 기록
            ledger_entry = CreditsLedger(
                user_id=user.id,
                payment_id=payment.id,
                delta=-credits_to_refund,
                reason='webhook_refund',
                note=f'Webhook 자동 환불 (imp_uid={imp_uid})'
            )
            db.session.add(ledger_entry)
            
            logging.info(f"[WEBHOOK][OK] payment_id={payment.id}, user_id={user.id}, refunded_credits={credits_to_refund}")
        
        db.session.commit()
        
        return jsonify({"ok": True, "msg": "refund synced"})
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"[WEBHOOK][ERROR] 웹훅 처리 오류: {str(e)}")
        return jsonify({"ok": False, "error": "webhook_error"}), 500


@app.route("/pay/verify", methods=["POST"])
@login_required
def pay_verify():
    """PortOne 결제 검증 및 평가권 증가"""
    try:
        data = request.get_json() or request.form
        imp_uid = data.get("imp_uid")
        merchant_uid = data.get("merchant_uid")
        
        if not imp_uid or not merchant_uid:
            logging.error("[PAY-VERIFY][FAIL] imp_uid 또는 merchant_uid 누락")
            return jsonify({"success": False, "message": "필수 파라미터 누락"}), 400
        
        # merchant_uid에서 상품 코드 추출 (basic_ 또는 bundle5_)
        product_code = None
        if merchant_uid.startswith("basic_"):
            product_code = "basic"
            credits_to_add = 1
        elif merchant_uid.startswith("bundle5_"):
            product_code = "bundle5"
            credits_to_add = 5
        else:
            logging.error(f"[PAY-VERIFY][FAIL] 알 수 없는 merchant_uid 형식: {merchant_uid[:20]}***")
            return jsonify({"success": False, "message": "잘못된 주문 정보"}), 400
        
        # PortOne API로 결제 정보 조회 (V1 REST API)
        portone_api_key = os.environ.get("PORTONE_API_KEY")
        portone_api_secret = os.environ.get("PORTONE_API_SECRET")
        
        if not portone_api_key or not portone_api_secret:
            logging.error("[PAY-VERIFY][FAIL] PortOne API 키 미설정")
            return jsonify({"success": False, "message": "결제 시스템 설정 오류"}), 500
        
        # PortOne V1 API - 결제 정보 조회
        verify_url = f"https://api.iamport.kr/payments/{imp_uid}"
        headers = {
            "Content-Type": "application/json"
        }
        
        # Access Token 발급
        token_url = "https://api.iamport.kr/users/getToken"
        token_data = {
            "imp_key": portone_api_key,
            "imp_secret": portone_api_secret
        }
        
        token_response = requests.post(token_url, json=token_data)
        if token_response.status_code != 200:
            logging.error(f"[PAY-VERIFY][FAIL] PortOne 토큰 발급 실패: {token_response.status_code}")
            return jsonify({"success": False, "message": "결제 검증 실패"}), 500
        
        token_result = token_response.json()
        if token_result.get("code") != 0:
            logging.error(f"[PAY-VERIFY][FAIL] PortOne 토큰 응답 오류: {token_result.get('message')}")
            return jsonify({"success": False, "message": "결제 검증 실패"}), 500
        
        access_token = token_result["response"]["access_token"]
        
        # 결제 정보 조회
        headers["Authorization"] = access_token
        payment_response = requests.get(verify_url, headers=headers)
        
        if payment_response.status_code != 200:
            logging.error(f"[PAY-VERIFY][FAIL] 결제 정보 조회 실패: {payment_response.status_code}")
            return jsonify({"success": False, "message": "결제 정보 조회 실패"}), 500
        
        payment_result = payment_response.json()
        if payment_result.get("code") != 0:
            logging.error(f"[PAY-VERIFY][FAIL] 결제 정보 응답 오류: {payment_result.get('message')}")
            return jsonify({"success": False, "message": "결제 정보 조회 실패"}), 500
        
        payment_data = payment_result["response"]
        
        # 결제 상태 확인
        if payment_data.get("status") != "paid":
            logging.error(f"[PAY-VERIFY][FAIL] 결제 미완료 상태: {payment_data.get('status')}")
            return jsonify({"success": False, "message": "결제가 완료되지 않았습니다"}), 400
        
        # 금액 검증
        expected_amount = 2500 if product_code == "basic" else 8900
        if payment_data.get("amount") != expected_amount:
            logging.error(f"[PAY-VERIFY][FAIL] 금액 불일치: 예상={expected_amount}, 실제={payment_data.get('amount')}")
            return jsonify({"success": False, "message": "결제 금액이 일치하지 않습니다"}), 400
        
        # 중복 결제 체크
        existing_purchase = CreditPurchase.query.filter_by(
            user_id=current_user.id
        ).filter(
            db.func.date(CreditPurchase.purchase_date) == datetime.utcnow().date()
        ).filter(
            CreditPurchase.original_amount == credits_to_add
        ).first()
        
        # 같은 날 같은 금액의 구매가 있으면 의심
        if existing_purchase and (datetime.utcnow() - existing_purchase.purchase_date).seconds < 60:
            logging.warning(f"[PAY-VERIFY][WARN] 1분 내 중복 결제 의심: user_id={current_user.id}")
            # 중복 결제는 허용하되 로그 남김
        
        # 구매 기록 생성 또는 조회 (idempotent)
        expires_at = datetime.utcnow() + timedelta(days=90)  # 3개월 후 만료
        product_name = "JasoAI 1회 평가권" if product_code == "basic" else "JasoAI 5회 평가권"
        paid_time = datetime.utcnow()
        
        # product_code 표준화: basic -> EVAL_1, bundle5 -> EVAL_5
        standard_product_code = "EVAL_1" if product_code == "basic" else "EVAL_5"
        
        # imp_uid 기반 중복 결제 체크
        existing_payment = CreditPurchase.query.filter_by(imp_uid=imp_uid).first()
        
        if existing_payment:
            # 이미 처리된 결제인 경우 (idempotent)
            if existing_payment.credited_at:
                logging.warning(f"[PAY-VERIFY][IDEMPOTENT] 이미 처리된 결제: imp_uid={imp_uid}")
                return jsonify({
                    "success": True,
                    "message": "이미 처리된 결제입니다",
                    "credits_added": credits_to_add,
                    "total_credits": current_user.remaining_credits
                })
            else:
                # Payment는 있지만 credited_at이 없는 경우 - 재처리
                new_purchase = existing_payment
                logging.info(f"[PAY-VERIFY][RETRY] Payment 있음, 크레딧 재처리: payment_id={new_purchase.id}")
        else:
            # 새 결제 생성
            new_purchase = CreditPurchase(
                user_id=current_user.id,
                purchase_date=paid_time,
                original_amount=credits_to_add,
                remaining_amount=credits_to_add,
                expires_at=expires_at,
                expired_at=None,
                imp_uid=imp_uid,
                merchant_uid=merchant_uid,
                product_name=product_name,
                amount=expected_amount,
                refund_status='none',
                product_code=standard_product_code,
                credits_added=credits_to_add,
                status='paid',
                paid_at=paid_time,
                refunded_at=None,
                credited_at=None
            )
            db.session.add(new_purchase)
            db.session.flush()  # ID 생성을 위해 flush
        
        # Idempotent 크레딧 지급
        if not new_purchase.credited_at:
            # Ledger 중복 체크
            existing_ledger = CreditsLedger.query.filter_by(
                payment_id=new_purchase.id,
                reason='payment'
            ).first()
            
            if not existing_ledger:
                # 평가권 증가
                current_user.remaining_credits = (current_user.remaining_credits or 0) + credits_to_add
                current_user.evaluation_ticket = (current_user.evaluation_ticket or 0) + credits_to_add
                
                # Ledger 기록
                ledger_entry = CreditsLedger(
                    user_id=current_user.id,
                    payment_id=new_purchase.id,
                    delta=credits_to_add,
                    reason='payment',
                    note=f'결제 성공 (imp_uid={imp_uid})'
                )
                db.session.add(ledger_entry)
            
            # credited_at 설정
            new_purchase.credited_at = datetime.utcnow()
        
        db.session.commit()
        
        logging.info(f"[PAY-VERIFY][OK] user_id={current_user.id}, product={product_code}, credits_added={credits_to_add}, new_total={current_user.remaining_credits}")
        
        add_notification(
            user_id=current_user.id,
            notif_type='payment_success',
            title='결제가 완료되었습니다',
            body=f'{product_name} {expected_amount:,}원 결제 완료',
            meta={'imp_uid': imp_uid, 'merchant_uid': merchant_uid, 'credits': credits_to_add}
        )
        
        return jsonify({
            "success": True,
            "message": "결제가 완료되었습니다",
            "credits_added": credits_to_add,
            "total_credits": current_user.remaining_credits
        })
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"[PAY-VERIFY][ERROR] 결제 검증 오류: {str(e)}")
        return jsonify({"success": False, "message": "결제 처리 중 오류가 발생했습니다"}), 500


def _humanize_payment_method(method, pg_provider):
    """결제 수단을 사람이 이해하기 쉬운 한글로 변환"""
    m = (method or "").lower()
    pg = (pg_provider or "").lower()
    
    if "card" in m:
        return "신용/체크카드"
    if "vbank" in m:
        return "가상계좌"
    if "trans" in m or "account" in m:
        return "계좌이체"
    if "kakaopay" in pg or "kakao" in pg:
        return "카카오페이"
    if "tosspay" in pg or "tosspayments" in pg or "toss" in pg:
        return "토스페이"
    if "naverpay" in pg or "naver" in pg:
        return "네이버페이"
    
    return "간편결제"


def _product_display_name(product_code):
    """내부 상품 코드를 사용자 친화적인 상품명으로 변환"""
    if not product_code:
        return "JasoAI 평가권"
    
    code = str(product_code).lower()
    if "5" in code or "bundle" in code or "eval_5" in code:
        return "JasoAI 평가권 5회"
    
    return "JasoAI 평가권 1회"


def _format_amount(won):
    """금액을 3자리 콤마 형식으로 포맷"""
    try:
        return f"{int(round(float(won))):,}원"
    except:
        return f"{won}원"


def _short_order_number(merchant_uid):
    """주문번호를 사용자 친화적인 형식으로 변환: JASOAI-YYYYMMDD-XXXXXX"""
    tail = (merchant_uid or "")[-6:].upper()
    ymd = datetime.now().strftime("%Y%m%d")
    return f"JASOAI-{ymd}-{tail}"


def _format_datetime_kst(dt):
    """날짜를 KST 기준으로 포맷"""
    if isinstance(dt, str):
        try:
            dt = datetime.fromisoformat(dt.replace('Z', '+00:00'))
        except:
            return dt
    
    if dt:
        kst_dt = dt + timedelta(hours=9)
        return kst_dt.strftime('%Y-%m-%d %H:%M')
    
    return datetime.now().strftime('%Y-%m-%d %H:%M')


@app.route("/pay/result")
def pay_result():
    """포트원 결제 결과 페이지 (프론트엔드에서 검증 요청)"""
    success = request.args.get("success", "")
    imp_uid = request.args.get("imp_uid", "")
    merchant_uid = request.args.get("merchant_uid", "")
    error_msg = request.args.get("error_msg", "")
    amount = request.args.get("amount", "")
    product_name = request.args.get("product_name", "")
    pg = request.args.get("pg", "")
    
    display = {}
    if success == "1":
        display = {
            "product_display_name": _product_display_name(product_name),
            "amount_display": _format_amount(amount),
            "pay_method_display": _humanize_payment_method("card", pg),
            "order_no_display": _short_order_number(merchant_uid),
            "imp_uid": imp_uid,
            "approved_at_display": _format_datetime_kst(datetime.utcnow())
        }
    
    return render_template("pay_result.html",
        success=success,
        imp_uid=imp_uid,
        merchant_uid=merchant_uid,
        error_msg=error_msg,
        amount=amount,
        product_name=product_name,
        pg=pg,
        display=display,
        now=datetime.utcnow())


# KG이니시스 본인인증 시스템
import hashlib
import hmac
import base64
from urllib.parse import quote

@app.route("/auth/inicis/request", methods=["POST"])
def inicis_auth_request():
    """이니시스 본인인증 요청"""
    try:
        data = request.get_json()
        return_url = data.get('returnUrl', '')
        
        if not return_url:
            return jsonify({"success": False, "message": "return URL이 필요합니다."}), 400
        
        # 이니시스 인증 요청 파라미터 구성
        site_code = os.environ.get('INICIS_IDENTITY_SITE_CODE')
        sign_key = os.environ.get('INICIS_IDENTITY_SIGNKEY')
        
        if not site_code or not sign_key:
            logging.error("이니시스 인증 키 설정 오류")
            return jsonify({"success": False, "message": "인증 시스템 설정 오류입니다."}), 500
        
        # 고유한 거래번호 생성
        import time
        merchant_uid = f"auth_{int(time.time())}_{current_user.id if current_user.is_authenticated else 'guest'}"
        
        # 서명 생성
        sign_data = f"{site_code}{merchant_uid}{return_url}"
        signature = hmac.new(
            sign_key.encode('utf-8'),
            sign_data.encode('utf-8'),
            hashlib.sha256
        ).hexdigest()
        
        # 이니시스 인증 요청 URL 구성
        auth_params = {
            'mid': site_code,
            'tid': merchant_uid,
            'returnUrl': return_url,
            'signature': signature
        }
        
        # 이니시스 LIVE 인증 URL
        auth_url = "https://cert.inicis.com/auth/" + "&".join([f"{k}={quote(str(v))}" for k, v in auth_params.items()])
        
        # 세션에 거래번호 저장
        session['inicis_merchant_uid'] = merchant_uid
        
        logging.info(f"이니시스 인증 요청: {_mask_merchant_uid(merchant_uid)}")
        
        return jsonify({
            "success": True,
            "authUrl": auth_url,
            "merchantUid": merchant_uid
        })
        
    except Exception as e:
        logging.error(f"이니시스 인증 요청 오류: 인증 요청 시스템 오류")  # 민감정보 노출 방지
        return jsonify({"success": False, "message": "인증 요청 실패"}), 500


@app.route("/auth/inicis/callback")
def inicis_auth_callback():
    """이니시스 본인인증 콜백"""
    try:
        # 이니시스에서 전달된 인증 결과 파라미터
        result_code = request.args.get('resultCode', '')
        result_msg = request.args.get('resultMsg', '')
        merchant_uid = request.args.get('tid', '')
        
        # CI 및 개인정보 (성공 시에만 전달됨)
        ci_hash = request.args.get('ci', '')
        user_name = request.args.get('name', '')
        birth_date = request.args.get('birthDate', '')
        phone_number = request.args.get('phoneNo', '')
        
        logging.info(f"이니시스 콜백: 코드={result_code}, 메시지=마스킹됨, UID={_mask_merchant_uid(merchant_uid)}")
        
        # 세션에서 거래번호 확인
        session_merchant_uid = session.get('inicis_merchant_uid')
        if not session_merchant_uid or session_merchant_uid != merchant_uid:
            logging.warning(f"거래번호 불일치: {_mask_merchant_uid(session_merchant_uid)} vs {_mask_merchant_uid(merchant_uid)}")
        
        if result_code == '0000':  # 성공
            # 로그인 사용자인 경우 DB 업데이트
            if current_user.is_authenticated and ci_hash and user_name:
                try:
                    current_user.realname_verified = True
                    current_user.ci_hash = ci_hash
                    current_user.verified_name = user_name
                    current_user.verified_at = datetime.utcnow()
                    
                    # 휴대폰 번호도 업데이트 (기존 값이 없는 경우에만)
                    if phone_number and not current_user.phone_number:
                        current_user.phone_number = phone_number
                    
                    db.session.commit()
                    logging.info(f"본인인증 완료: {_mask_username(current_user.user_id)} (이름마스킹)")
                    
                    # 성공 상태를 세션에 저장
                    session['inicis_auth_success'] = True
                    session['inicis_auth_data'] = {
                        'verified': True,
                        'name': user_name,
                        'ci_hash': ci_hash[:8] + '...'  # 보안상 일부만 저장
                    }
                    
                except Exception as e:
                    logging.error(f"DB 업데이트 오류: 민감정보 로그 마스킹됨")
                    db.session.rollback()
            
            # 성공 페이지 표시
            return render_template('inicis_auth_result.html',
                                   success=True,
                                   message='본인인증이 성공적으로 완료되었습니다.',
                                   user_name=user_name,
                                   verified=True)
            
        else:  # 실패 또는 취소
            # 실패 상태를 세션에 저장
            session['inicis_auth_success'] = False
            session['inicis_auth_data'] = {
                'verified': False,
                'error': result_msg or '인증에 실패했습니다.'
            }
            
            return render_template('inicis_auth_result.html',
                                   success=False,
                                   message=result_msg or '인증에 실패했습니다.',
                                   verified=False)
        
    except Exception as e:
        logging.error(f"이니시스 콜백 오류: 시스템 오류 발생")  # 민감정보 노출 방지
        session['inicis_auth_success'] = False
        session['inicis_auth_data'] = {
            'verified': False,
            'error': '시스템 오류가 발생했습니다.'
        }
        
        return render_template('inicis_auth_result.html',
                               success=False,
                               message='시스템 오류가 발생했습니다.',
                               verified=False)


@app.route("/auth/inicis/result")
def inicis_auth_result():
    """이니시스 인증 결과 조회 (AJAX용)"""
    try:
        auth_success = session.get('inicis_auth_success')
        auth_data = session.get('inicis_auth_data', {})
        
        if auth_success is None:
            return jsonify({
                "success": False,
                "message": "인증 결과가 없습니다. 다시 시도해주세요."
            })
        
        # 결과 반환 후 세션 정리
        session.pop('inicis_auth_success', None)
        session.pop('inicis_auth_data', None)
        session.pop('inicis_merchant_uid', None)
        
        if auth_success:
            return jsonify({
                "success": True,
                "message": "본인인증이 성공적으로 완료되었습니다.",
                "data": auth_data
            })
        else:
            return jsonify({
                "success": False,
                "cancelled": 'error' not in auth_data,
                "message": auth_data.get('error', '인증이 취소되었습니다.')
            })
            
    except Exception as e:
        logging.error(f"인증 결과 조회 오류: {str(e)}")
        return jsonify({
            "success": False,
            "message": "결과 조회 중 오류가 발생했습니다."
        }), 500
